#!/usr/local/bin/python3 -u

import ezdxf
import sys, os
from pprint import pprint
from ezdxf import bbox
from ezdxf.addons import Importer
from ezdxf.math import Vec2, intersection_line_line_2d, convex_hull_2d
from datetime import date

import openpyxl
from openpyxl.styles import PatternFill
from openpyxl.styles import Alignment
import docx
import queue
import bisect
import threading
import numpy as np
from openpyxl.styles.borders import Border, Side
import os.path
from math import ceil, floor, sqrt, log10, atan2, pi
from copy import copy, deepcopy
from tkinter import *
from tkinter import filedialog
from tkinter.messagebox import askyesno

dxf_version = "AC1032"

web_version = False
debug = False
frame_enabled = False

if ezdxf.version == (0, 14, 2, 'release'):
    poly_class = ezdxf.entities.lwpolyline.LWPolyline
    poly_class.is_closed = poly_class.close

if len(sys.argv) > 1 :
	web_version = True

# block names (defaults)
block_blue_120x100  = "LEO_55_120"
block_blue_60x100   = "LEO_55_60"
block_green_120x100 = "LEO_55_120_IDRO"
block_green_60x100  = "LEO_55_60_IDRO"
block_collector     = "collettore"

# Panel types
panel_types = [
    {
        "full_name"  : "Leonardo 5,5",
        "handler"    : "55",
        "rings"      : 10,
        "panels"     : 5,
        "flow_line"  : 280,
        "flow_ring"  : 28,
        "flow_panel" : 56,
		"code_full"  : "6113010431",
		"desc_full"  : "LEONARDO 5,5 MS - 1200x2000x50mm",
		"code_half"  : "6113010432",
		"desc_half"  : "LEONARDO 5,5 MS - 600x2000x50mm",
		"code_full_h": "6114010411",
		"desc_full_h": "LEONARDO 5,5 IDRO MS - 1200x2000x50mm",
		"code_half_h": "6114010412",
		"desc_half_h": "LEONARDO 5,5 IDRO MS - 600x2000x50mm",
    },
    {
        "full_name"  : "Leonardo 3,5",
        "handler"    : "35",
        "rings"      : 9,
        "panels"     : 4.5,
        "flow_line"  : 252,
        "flow_ring"  : 28,
        "flow_panel" : 56,
		"code_full"  : "6113010451",
		"desc_full"  : "LEONARDO 3,5 MS - 1200x2000x50mm",
		"code_half"  : "6113010452",
		"desc_half"  : "LEONARDO 3,5 MS - 600x2000x50mm",
		"code_full_h": "6114010431",
		"desc_full_h": "LEONARDO 3,5 IDRO MS - 1200x2000x50mm",
		"code_half_h": "6114010432",
		"desc_half_h": "LEONARDO 3,5 IDRO MS - 600x2000x50mm",
    },
    {
        "full_name"  : "Leonardo 3,0 plus",
        "handler"    : "30",
        "rings"      : 9,
        "panels"     : 4.5,
        "flow_line"  : 265,
        "flow_ring"  : 29.4,
        "flow_panel" : 58.9,
		"code_full"  : "6113011001",
		"desc_full"  : "LEONARDO 3,0 PLUS MS - 1200x2000x50mm",
		"code_half"  : "6113011002",
		"desc_half"  : "LEONARDO 3,0 PLUS MS - 600x2000x50mm",
		"code_full_h": "6113011001",
		"desc_full_h": "LEONARDO 3,0 PLUS MS - 1200x2000x50mm",
		"code_half_h": "6113011002",
		"desc_half_h": "LEONARDO 3,0 PLUS MS - 600x2000x50mm",
    }
]

air_conditioners = [
	{
		"type": "dehum",
		"type_label": "Deumidificatore",
		"model": "DEUMIDIFICATORE 581 DC",
		"mount": "O",
		"width_mm": [756],
		"height_mm": [260],
		"depth_mm": [803],
		"flow_m3h": 300,
		"code": "7110010301",
		"accessories": {"sifone", "plenum"}
	},
	{
		"type": "dehum",
		"type_label": "Deumidificatore",
		"model": "DEUMIDIFICATORE 901 DC",
		"mount": "O",
		"width_mm": [706],
		"height_mm": [309],
		"depth_mm": [936],
		"flow_m3h": 580,
		"code": "7110010601",
		"accessories": {"sifone"}
	},
	{
		"type": "dehum",
		"type_label": "Deumidificatore",
		"model": "DEUMIDIFICATORE 320 DI (incasso)",
		"mount": "V",
		"width_mm": [402],
		"height_mm": [637],
		"depth_mm": [203],
		"flow_m3h": 200,
		"code": "7110020101",
		"accessories": {"sifone", "telaio1", "griglia1"}
	},
	{
		"type": "dehum",
		"type_label": "Deumidificatore",
		"model": "DEUMIDIFICATORE 581 DI (incasso)",
		"mount": "V",
		"width_mm": [732],
		"height_mm": [732],
		"depth_mm": [203],
		"flow_m3h": 300,
		"code": "7110020101",
		"accessories": {"sifone", "telaio2", "griglia2"}
	},
	{
		"type": "dehum_int",
		"type_label": "Deuclimatizzatore",
		"model": "DEU-CLIMATIZZATORE 582 DCC",
		"mount": "O",
		"width_mm": [756],
		"depth_mm": [260],
		"height_mm": [803],
		"flow_m3h": 300,
		"code": "7210020701",
		"accessories": {"sifone", "plenum"}
	},
	{
		"type": "dehum_int",
		"type_label": "Deuclimatizzatore",
		"model": "DEU-CLIMATIZZATORE 901 DCC",
		"mount": "O",
		"width_mm": [706],
		"height_mm": [309],
		"depth_mm": [936],
		"flow_m3h": 580,
		"code": "7210010602",
		"accessories": {"sifone"}
	},
	{
		"type": "dehum_int",
		"type_label": "Deuclimatizzatore",
		"model": "DEU-CLIMATIZZATORE 581 DCI (incasso)",
		"mount": "V",
		"width_mm": [732],
		"height_mm": [732],
		"depth_mm": [203],
		"flow_m3h": 300,
		"code": "7210020301",
		"accessories": {"sifone", "telaio2", "griglia2"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEUCLIMA-VMC 300S",
		"mount": "O",
		"width_mm": [1204.4],
		"height_mm": [979.5],
		"depth_mm": [244],
		"flow_m3h": 300,
		"code": "7410010103",
		"accessories": {("sifone",2), "filtro"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEUCLIMA-VMC 500S",
		"mount": "O",
		"width_mm": [1254.4],
		"height_mm": [810.5],
		"depth_mm": [294],
		"flow_m3h": 500,
		"code": "7410010105",
		"accessories": {("sifone",2), "filtro"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEUCLIMA-VMC 300V",
		"mount": "V",
		"width_mm": [1391.7],
		"height_mm": [700],
		"depth_mm": [342.3],
		"flow_m3h": 300,
		"code": "7510010101",
		"accessories": {"sifone", "filtro"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEUCLIMA-VMC 500V",
		"mount": "V",
		"width_mm": [1696.7],
		"height_mm": [700],
		"depth_mm": [421],
		"flow_m3h": 500,
		"code": "7510010102",
		"accessories": {"sifone", "filtro"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEU-CLIMATIZZATORE DCR 1000",
		"mount": "O",
		"width_mm": [805,1097],
		"height_mm": [691,723],
		"depth_mm": [350.5,350.5],
		"flow_m3h": 1000,
		"code": "7110011001",
		"accessories": {("sifone",2), "dcr1000"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEU-CLIMATIZZATORE DCR 2000",
		"mount": "O",
		"width_mm": [950.5,1097],
		"height_mm": [691,723],
		"depth_mm": [350.5,350.5],
		"flow_m3h": 2000,
		"code": "7110011002",
		"accessories": {("sifone", 2), "dcr2000"}
	},
	{
		"type": "dehum_int_ren",
		"type_label": "Deuclima VMC",
		"model": "DEUCLIMA-VMC 300SY",
		"mount": "O",
		"width_mm": [1070],
		"height_mm": [880],
		"depth_mm": [251],
		"flow_m3h": 300,
		"code": "7410010104",
		"accessories": {"sifone"}
	},
]

accessories = {
	"sifone": {
		"code": "7910080907",
		"desc": "Sifone"
	},
	"plenum": {
		"code": "7510030203",
		"desc": "Plenum di mandata 2/3xD160"
	},
	"telaio1": {
		"code": "7910010201",
		"desc": "Telaio da murare"
	},
	"griglia1": {
		"code": "7910010101",
		"desc": "Griglia in legno laccato"
	},
	"telaio2": {
		"code": "7910030201",
		"desc": "Telaio da murare"
	},
	"griglia2": {
		"code": "7910030101",
		"desc": "Griglia in legno laccato"
	},
	"filtro": {
		"code": "7910080980",
		"desc": "Box filtro con lampada UV e silenziatore"	
	},
	"dcr1000": {
		"code": "7910090901",
		"desc": "Silenziatore DCR1000"
	},
	"dcr2000": {
		"code": "7910090902",
		"desc": "Silenziatore DCR2000"
	},
}


ac_label = {
	"dehum":         "Deumidificatore",
	"dehum_int":     "Deuclimatizzatore",
	"dehum_int_ren": "Deuclima VMC"
}


# couplings geometry
stripe_width = 0.08
stripe_offset = 0.44
delta_h  = 2.5
delta_v  = 4
axis_offset = 12
shift = 13

w_fit = 10
h_fit = 5
s_fit = 7

# frame geometry
a0_ax = -74680
a0_ay = -78829
a0_bx = -15227
a0_by = -36723

a1_ax = -10444
a1_ay = -79890
a1_bx =  31439
a1_by = -49108

fittings = {
	"Rac_20_10_20": {
		"desc":      "20-10-20",
		"code":      "6910022007",
		"symbol": [ 
					[(0,0),(0,h_fit)], 
					[(-w_fit,0),(w_fit,0)]]
	},
	"Rac_20_10": {
		"desc":      "20-10",
		"code":      "6910022107",
		"symbol": [
					[(0,0),(0,h_fit)], 
					[(0,0),(w_fit,0)]]
	},
	"Rac_20_10_20_10": {
		"desc":      "20-10-20-10",
		"code":      "6910022003",
		"symbol": [
					[(0,-h_fit),(0,h_fit)], 
					[(-w_fit,0),(w_fit,0)]]
	},
	"Rac_10_20_10": {
		"desc":      "10-20-10",
		"code":      "6910022103",
		"symbol": [
					[(0,-h_fit),(0,h_fit)], 
					[(0,0),(w_fit,0)]]
	},
	"Rac_20_10_10_20": {
		"desc":      "20-10-10-20",
		"code":      "6910022008",
		"symbol": [
					[(s_fit,0),(s_fit,h_fit)], 
					[(-w_fit,0),(w_fit,0)], 
					[(-s_fit,0),(-s_fit,h_fit)]]
	},
	"Rac_20_10_10": {
		"desc":      "20-10-10",
		"code":      "6910022108",
		"symbol": [
					[(s_fit,0),(s_fit,h_fit)], 
					[(-s_fit,0),(w_fit,0)], 
					[(-s_fit,0),(-s_fit,h_fit)]]
	},
	"Rac_10_20_10_10": {
		"desc":      "10-20-10-10",
		"code":      "6910022110",
		"symbol": [
					[(s_fit,0),(s_fit,h_fit)], 
					[(-s_fit,0),(w_fit,0)], 
					[(-s_fit,0),(-s_fit,h_fit)]]
	},
	"Rac_20_10_10_20_10_10": {
		"desc":      "20-10-10-20-10-10",
		"code":      "6910022004",
		"symbol": [
					[(s_fit,-h_fit),(s_fit,h_fit)], 
					[(-w_fit,0),(w_fit,0)], 
					[(-s_fit,-h_fit),(-s_fit,h_fit)]]
	},
	"Rac_10_10_20_10_10": {
		"desc":      "10-10-20-10-10",
		"code":      "6910022104",
		"symbol":   [
					[(s_fit,-h_fit),(s_fit,h_fit)], 
					[(-s_fit,0),(w_fit,0)], 
					[(-s_fit,0),(-s_fit,h_fit)]]
	},
	"Rac_20_20_dritto": {
		"desc":      "20-20 Dritto",
		"code":      "6910022005",
		"symbol":   [[(s_fit,-h_fit),(s_fit,h_fit)]] 
	},
	"Rac_20_20_curva": {
		"desc":      "20-20 Curva",
		"code":      "6910022006",
		"symbol":   [[(s_fit,-h_fit),(s_fit,h_fit)]] 
	},
	"Rac_20_20_20": {
		"desc":      "20-20-20 (4pz)",
		"code":      "6910022009",
		"symbol":   [[(s_fit,-h_fit),(s_fit,h_fit)]] 
	},
	"Rac_10_10": {
		"desc":      "10-10",
		"code":      "6910022013",
		"symbol":   [[(s_fit,-h_fit),(s_fit,h_fit)]] 
	}
}


#names = [panel['full_name'] for panel in panel_types]
#print(names)
#exit()

# Parameter settings (values in cm)
default_scale = 1  # scale=100 if the drawing is in m
default_tolerance    = 1   # ignore too little variations

extra_len    = 20
zone_cost = 1
min_room_area = 1
max_room_area = 500
default_max_clt_distance = 3000
max_clt_break = 5

feeds_per_collector = 13
area_per_feed_m2 = 12
flow_per_m2 = 23.33
target_eff = 0.7
flow_per_collector = 1700

default_font_size = 10

# Half panels default dimensions in cm
default_panel_width = 100
default_panel_height = 60
default_hatch_width = 12
default_hatch_height = 20
default_collector_size = 60

default_search_tol = 5
default_min_dist = 20
default_min_dist2 = default_min_dist*default_min_dist
default_wall_depth = 50

# pipes geometry in cm
default_add_offs = 10
default_add_dist = 4

default_input_layer = 'AREE LEONARDO'
layer_text      = 'Eurotherm_text'
layer_box       = 'Eurotherm_box'
layer_panel     = 'Pannelli Leonardo'
layer_panelp    = 'Eurotherm_prof'
layer_link      = 'Eurotherm_link'
layer_error     = 'Eurotherm_error'
layer_lux       = 'Eurotherm_lux'
layer_probes    = 'Eurotherm_probes'
layer_joints    = 'Eurotherm_joints'
layer_struct    = 'Eurotherm_structure'
layer_collector = 'Collettori'

text_color = 7
box_color = 8               ;# cyan
collector_color = 1         ;# red
disabled_room_color = 6     ;# magenta
valid_room_color = 5        ;# blue
bathroom_color = 3          ;# green
obstacle_color = 2          ;# yellow
stripe_color = 1            ;# red
zone_color = 4              ;# cyan
color_warm = 1
color_cold = 5

ask_for_write = False

MAX_COST  = 1000000
MAX_DIST  = 1e20
MAX_DIST2 = 1e20

RIGHT = 1
LEFT = 0
TOP = 1
BOTTOM = 0

xlsx_template = 'leo_template.xlsx'
sheet_template_1 = 'LEONARDO 5.5'
sheet_template_2 = 'LEONARDO 3.5'
sheet_template_3 = 'LEONARDO 3.0 PLUS'

sheet_breakdown = [
	('Dettaglio Stanze L55', 85.0, 51.8), 
	('Dettaglio Stanze L35', 84.2, 64.8),
	('Dettaglio Stanze 30p', 82.3, 80.9)
]

#show_panel_list = True

alphabet = {
	' ': [],
	'A': [12,3,1,5,14,8,6],
	'B': [0,1,5,7,6,7,11,13,12,0],
	'C': [2,1,3,9,13,14],
	'D': [0,1,5,11,13,12,0],
	'E': [2,0,6,7,6,12,14],
	'F': [2,0,6,7,6,12],
	'G': [2,1,3,9,13,14,8,7],
	'H': [0,12,6,8,2,14],
	'I': [1,13],
	'J': [9,13,1,0,1,2],
	'K': [0,6,2,6,12,6,14],
	'L': [0,12,14],
	'M': [12,0,7,2,14],
	'N': [12,0,14,2],
	'O': [3,1,5,11,13,9,3],
	'P': [12,0,1,5,7,6],
	'Q': [13,11,5,1,3,9,13,14,13,10],
	'R': [12,0,1,5,7,6,7,11,14],
	'S': [12,13,11,3,1,2],
	'T': [0,2,1,13],
	'U': [0,9,13,11,2],
	'V': [0,13,2],
	'W': [0,13,2],
	'X': [0,14,7,2,12],
	'Y': [0,7,2,7,13],
	'Z': [0,2,12,14],
	'0': [1,3,9,13,11,5,1],
	'1': [3,1,13,12,14],
	'2': [3,1,5,12,14],
	'3': [0,1,5,7,6,7,11,13,12],
	'4': [0,6,8,2,14],
	'5': [2,0,6,7,11,13,12],
	'6': [2,1,3,9,13,11,7,6],
	'7': [0,2,12],
	'8': [3,1,5,9,12,14,11,3],
	'9': [8,6,3,1,5,11,13,12]
}




def crp(x, y):
	return x[0]*y[0] + x[1]*y[1]


def nav_item(qnty, code, desc):

	if qnty<=0:
		return ""

	if (type(qnty)==int):
		txt = '%d\r\n' % qnty
	else:
		txt = ('%.1f\r\n' % qnty).replace(".", ",")
	
	txt += code
	txt += '      EURO\t'
	txt += desc + '\r\n'

	return txt


max_iterations = 5e6

def pletter(letter,pos,scale):
	path = alphabet[letter]
	poly = []
	for p in path:
		x = p % 3
		y = 4 - (p//3)
		poly.append((pos[0]+x*scale[0], pos[1]+y*scale[1]))
	return poly

def writedxf(msp, strn, pos, scale):

	slen = len(strn)*scale[0]*3
	pos = (pos[0]-slen/2, pos[1]-scale[1]*2)
	
	for l in strn:
		poly = pletter(l, pos, scale)
		pl = msp.add_lwpolyline(poly)
		pl.dxf.layer = layer_text
		pos = (pos[0]+scale[0]*3, pos[1])

def write_text(msp, strn, pos, 
	align=ezdxf.lldxf.const.MTEXT_MIDDLE_CENTER, zoom=1, col=text_color):
	
	text = msp.add_mtext(strn, 
		dxfattribs={"style": "Arial"})
	text.dxf.insert = pos
	text.dxf.attachment_point = align
	text.dxf.char_height = font_size*zoom
	text.dxf.layer = layer_text
	text.dxf.color = col
	

def spread(a,b,size):
	l = []
	n = ceil(abs(a-b)/size) - 1
	L = abs(a-b)/(n+1)
	for i in range(0,n):
		l.append((i+1)*L+a)
	return l

def center(a,b,size):
	l = []
	n = ceil(abs(b-a)/size) - 1
	L = (abs(b-a) - size*n)/2
	for i in range(0,n+1):
		l.append(a+L+i*size)	
	return l

def dist(point1, point2):
	dx = point1[0] - point2[0]
	dy = point1[1] - point2[1]
	return sqrt(dx*dx+dy*dy)
	

def intersects(p, x):
		
	ints = []
	eps = 1e-6
		
	for i in range(0, len(p)-1):
		vx, vy = p[i+1][0]-p[i][0], p[i+1][1]-p[i][1]

		if (p[i+1][0]==x or p[i][0]==x):
				x += eps

		if (p[i][0] < p[i+1][0]):
			x0, y0 = p[i][0], p[i][1]
			x1, y1 = p[i+1][0], p[i+1][1]
		else:
			x0, y0 = p[i+1][0], p[i+1][1]
			x1, y1 = p[i][0], p[i][1]
					
		if (x0<x and x<x1):
			delta = x1 - x0
			py = (y0*(x1-x) - y1*(x0-x))/delta
			ints.append(py)
	
	return sorted(ints)


# Check if line crosses box
def cross(box, line):

	p0, p1 = line[0], line[1]
	p0x = p0[0]; p0y = p0[1]
	p1x = p1[0]; p1y = p1[1]
	ax = box[0]; bx = box[1]
	ay = box[2]; by = box[3]

	# both left/right/up/down
	if ( (p0x<=ax and p1x<=ax) or
	     (p0x>=bx and p1x>=bx) or
		 (p0y<=ay and p1y<=ay) or
		 (p0y>=by and p1y>=by)):
		return False

	# vertical or horizontal line inside
	if (p0x == p1x or p0y == p1y):
		return True

	# lateral crossing 
	A1 = (ay - p0y) * (p1x - p0x)
	A2 = (ax - p0x) * (p1y - p0y)
	A3 = (by - p0y) * (p1x - p0x)
	A4 = (bx - p0x) * (p1y - p0y)

	if (p1x>p0x):
		if (p1y>p0y):
			return  ((A1<=A2 and A2<=A3) or
				     (A2<=A1 and A1<=A4) or
				     (A1<=A4 and A4<=A3) or
				     (A2<=A3 and A3<=A4))
		else:
			return  ((A1<=A2 and A2<=A3) or
				     (A2>=A1 and A1>=A4) or
				     (A1<=A4 and A4<=A3) or
				     (A2>=A3 and A3>=A4))
	else:
		if (p1y>p0y):
			return  ((A1>=A2 and A2>=A3) or
				     (A2<=A1 and A1<=A4) or
				     (A1>=A4 and A4>=A3) or
				     (A2<=A3 and A3<=A4))
			
		else:
			return  ((A1>=A2 and A2>=A3) or
				     (A2>=A1 and A1>=A4) or
				     (A1>=A4 and A4>=A3) or
				     (A2>=A3 and A3>=A4))
	return False


def hdist(line, point):
	(xp,yp) = point
	(x0,y0),(x1,y1) = line
	
	(x0,y0) = (x0-xp,y0-yp)
	(x1,y1) = (x1-xp,y1-yp)

	if ((y0>0 and y1>0) or (y0<0 and y1<0)):
		return MAX_DIST

	if (y0==y1):
		return min(x1,x0)

	return abs((x0*y1-x1*y0)/(y1-y0))


def dist2(line, point):
	(xp,yp) = point
	(x0,y0),(x1,y1) = line
	
	(ux, uy) = (x1-x0, y1-y0)
	(px, py) = (xp-x0, yp-y0)
	
	u2 = ux*ux + uy*uy
	p2 = px*px + py*py
	up = ux*px + uy*py

	if (u2==0):
		return (0, False, point)

	q = (x0 + up*ux/u2, y0 + up*uy/u2)

	if (up>u2):
		return (u2+p2-2*up, False, line[1])

	if (up<=0 or u2==0):
		return (p2, False, line[0])

	return (p2-up*up/u2, True, q)



# Project line1 into line2 and returns the 
# facing segments
def is_gate(line, target):
	
	(xa0,ya0), (xa1,ya1) = (a0,a1) = line
	(xb0,yb0), (xb1,yb1) = (b0,b1) = target

	# reference on target
	(ux, uy) = (xb1-xb0, yb1-yb0)
	u = sqrt(ux*ux + uy*uy)
	if (u==0):
		return (False, (None, None))
	(uvx, uvy) = (ux/u, uy/u)
	(uox, uoy) = (-uvy, uvx)
	
	# change of reference for p and q
	(px, py) = (xa0-xb0, ya0-yb0)
	(qx, qy) = (xa1-xb0, ya1-yb0)
	(npx, npy) = (uvx*px+uvy*py, uox*px+uoy*py)
	(nqx, nqy) = (uvx*qx+uvy*qy, uox*qx+uoy*qy)

	w = abs(npx - nqx)

	if ((w <= min_dist) or
		(npx<=0 and nqx<=0) or (npx>=u and nqx>=u)):
		return (False, (None, None))

	if (npx<0):
		p1 = (xb0,yb0)
		l1 = 0
	else:
		if (npx>u):
			p1 = (xb1,yb1)
			l1 = u
		else:
			p1 = (xb0+uvx*npx, yb0+uvy*npx)
			l1 = npx

	if (nqx<0):
		p2 = (xb0,yb0)
		l2 = 0
	else:
		if (nqx>u):
			p2 = (xb1,yb1)
			l2 = u
		else:
			p2 = (xb0+uvx*nqx, yb0+uvy*nqx)
			l2 = nqx

	if (abs(l1-l2)<=min_dist):
		return (False, (None, None))

	d1 = abs((npy*(nqx-l1) - (npx-l1)*nqy)/w)
	d2 = abs((npy*(nqx-l2) - (npx-l2)*nqy)/w)
	
	if (d1<= wall_depth and d2<=wall_depth):
		return (True, (p1, p2))

	return (False, (None, None))



def rotate(point, orig, uv):

	(ax, ay) = (orig[0] - point[0], orig[1] - point[1])
	bx =  ax*uv[0] + ay*uv[1]
	by = -ax*uv[1] + ay*uv[0]
	return (orig[0]+bx , orig[1]+by)

def set_border(ws, row, cols):

#	thin_border = Border(left=Side(style='thin'), 
#                   right=Side(style='thin'), 
#                   top=Side(style='thin'), 
#                   bottom=Side(style='thin'))

	cell_border = Border(top=Side(style='thin'))
	for col in cols:
		ws[col+row].border = cell_border

def project_hor(point, poly):

	if len(poly)<1:
		return point

	x = point[0]
	y = point[1]

	xmin = xmax = poly[0][0]
	ymin = ymax = poly[0][1]

	for i in range(len(poly)-1):
		pa = poly[i]
		pb = poly[i+1]
		if pa[1] == pb[1]:
			if pa[1] == y:
				return pa
		else:
			if pa[1] < pb[1]:
				if pa[1]<=y and y<=pb[1]:
					u = pb[0]*(y-pa[1]) + pa[0]*(pb[1]-y)
					d = pb[1] - pa[1]
					return (u/d, y)	
			else:
				if pb[1]<=y and y<=pa[1]:
					u = pa[0]*(y-pb[1]) + pb[0]*(pa[1]-y)
					d = pa[1] - pb[1]
					return (u/d, y)

		if ymin > pb[1]:
			xmin = pb[0]
			ymin = pb[1]
		
		if ymax < pb[1]:
			xmin = pb[0]
			ymax = pb[1]

	if y<ymin:
		return (xmin, y)

	return (xmax, y)
	

def project_vect(point, v, poly):
	return poly[0]


def miter(poly, offs):

	if len(poly)<3:
		return deepcopy(poly)
	
	opoly = [poly[0]]
	
	for i in range(1, len(poly)-1):
		u1x = poly[i-1][0] - poly[i][0]
		u1y = poly[i-1][1] - poly[i][1]
		d1u = sqrt(u1x*u1x + u1y*u1y)
		u2x = poly[i+1][0] - poly[i][0]
		u2y = poly[i+1][1] - poly[i][1]
		d2u = sqrt(u2x*u2x + u2y*u2y)
		
		if (d1u==0 or d2u==0):
			opoly.append(poly[i])
			continue

		off = min(offs, d1u/3, d2u/3)
	
		u1 = off*u1x/d1u, off*u1y/d1u
		u2 = off*u2x/d2u, off*u2y/d2u
		
		o1 = poly[i][0] + u1[0], poly[i][1] + u1[1]
		o2 = poly[i][0] + u2[0], poly[i][1] + u2[1]
		opoly.append(o1)
		opoly.append(o2)

	opoly.append(poly[-1])
	return opoly

def offset(poly, off):

	if len(poly)<2:
		return []

	opoly = list()

	up = u = (0, 0)
	for i in range(len(poly)-1):
		ux = poly[i+1][0] - poly[i][0]
		uy = poly[i+1][1] - poly[i][1]
		du = ux*ux + uy*uy
		if du == 0:
			continue
		du = sqrt(du)

		u = -uy/du, ux/du
		k = 1 + u[0]*up[0] + u[1]*up[1]
		if  k != 0:
			ou = (u[0]+up[0])/k, (u[1]+up[1])/k

		p = poly[i][0]+off*ou[0], poly[i][1]+off*ou[1]
		opoly.append(p)

		up = u

	opoly.append( (poly[-1][0]+off*u[0], poly[-1][1]+off*u[1]) )

	return opoly


def line_cross(line0, line1):

	a0x = line0[0][0]
	a0y = line0[0][1]
	a1x = line0[1][0]
	a1y = line0[1][1]
	b0x = line1[0][0]
	b0y = line1[0][1]
	b1x = line1[1][0]
	b1y = line1[1][1]

	delta  = -(a1x-a0x)*(b1y-b0y)+(b1x-b0x)*(a1y-a0y)
	if delta==0:
		return None

	deltat = -(b0x-a0x)*(b1y-b0y)+(b1x-b0x)*(b0y-a0y)
	deltas =  (a1x-a0x)*(b0y-a0y)-(a1y-a0y)*(b0x-a0x)
	t = deltat/delta
	s = deltas/delta

	if not (0<=s and s<=1 and 0<=t and t<=1):
		return None

	return (b0x + s*(b1x-b0x), b0y + s*(b1y-b0y))

def extend_to_poly(v, poly):

	ux = v[1][0] - v[0][0]
	uy = v[1][1] - v[0][1]

	if ux==0 and uy==0:
		return None

	# find BBox
	minx = maxx = poly[0][0]
	miny = maxy = poly[0][1]
	for i in range(1, len(poly)-1):
		minx = min(minx, poly[i][0])
		maxx = max(maxx, poly[i][0])
		miny = min(miny, poly[i][1])
		maxy = max(maxy, poly[i][1])

	if ((maxx < v[0][0] and ux>0) or 
		(minx > v[0][0] and ux<0) or
		(maxy < v[0][1] and uy>0) or 
		(miny > v[0][1] and uy<0)):
		return None

	fx = fy = MAX_DIST
	if not ux==0:
		if ux>0:
			dx = maxx - v[0][0]
		else:
			dx = minx - v[0][0]
		fx = dx/ux

	if not uy==0:
		if uy>0:
			dy = maxy - v[0][1]
		else:
			dy = miny - v[0][1]
		fy = dy/uy

	f = 2*min(fx, fy)
	p = (v[0][0] + f*ux, v[0][1]+f*uy)

	crss = list()
	l1 = (v[0], p)
	for i in range(len(poly)-1):
		pa = poly[i]
		pb = poly[i+1]
		l0 = (pa, pb)
		r = line_cross(l0, l1)
		if not r:
			continue
		crss.append(r)

	if len(crss) == 0:
		return None

	norm = dist(v[1], v[0])
	e = None
	d = MAX_DIST
	for c in crss:
		d1 = dist(c, v[0])/norm
		if d1>=0.99 and d1<d:
			d = d1
			e = c

	return e


def trim(poly, seg):
	
	crss = list()
	for i in range(len(poly)-1):
		line = (poly[i], poly[i+1])
		r = line_cross(line, seg)
		if r:
			d = dist(seg[0], r)
			crss.append([d, r])

	l = len(crss)
	if not (l>0 and l%2==0):
		return []

	crss.sort(key=lambda x: x[0])

	cr = list()
	for i in range(0,l,2):
		cr.append([crss[i][1], crss[i+1][1]])

	return cr


class Word:

	@classmethod
	def getv(cls, val):
		num = "{:0,.2f}".format(val)
		num = num.replace(',','$')
		num = num.replace('.',',')
		num = num.replace('$','.')
		return num

	@classmethod
	def euro(cls, val):
		eur = '\u20AC '
		num = cls.getv(val)
		return eur + num

	@classmethod
	def date(cls):
		today = date.today()
		str_date = str(today.day) 
		str_date +=  '/' + str(today.month) 
		str_date +=  '/' + str(today.year)
		return str_date

	def __init__(self, doc):
		self.document = docx.Document(doc)

	def write(self, para, text):
		self.document.paragraphs[para].runs[0].text = text

	def writetab(self, tabnum, cell, text):
		tab = self.document.tables[tabnum]
		tab.cell(*cell).paragraphs[0].runs[0].text = text

	def gettab(self, tabnum, cell):
		tab = self.document.tables[tabnum]
		num = tab.cell(*cell).paragraphs[0].runs[0].text

		num = num.replace('.','')
		num = num.replace(',','.')
		
		return float(num)


# This class represents the radiating panel
# with its characteristics
# side=00  -->  water feed over top edge, left panel
# side=01  -->  water feed over bottom edge, left panel
# side=10  -->  water feed over top edge, right panel
# side=11  -->  water feed over bottom edge, right panel
class Panel:
	def __init__(self, cell, size, side):
		self.side = side
		self.cell = cell
		self.xcoord = self.cell.box[0]
		self.ycoord = self.cell.box[2]
		self.width  = (self.cell.box[1] - self.xcoord) * size[0]
		self.height = (self.cell.box[3] - self.ycoord) * size[1]
		self.size = size
		self.color = cell.color
		self.mode = cell.mode
		self.pos = (self.xcoord, self.ycoord)
		self.stripe = False

	def center(self):
		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = ay + self.height

		cx = (ax+bx)/2
		cy = (ay+by)/2
		
		if (self.mode==1):
			cx, cy = cy, cx

		c = cx, cy

		if (self.cell.room.vector):
			room = self.cell.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			c = rotate(c, rot, uv)

		return c

	def polyline(self):

		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = ay + self.height

		pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]
		
		if (self.mode==1):
			for i in range(0,len(pline)):
				pline[i] = (pline[i][1], pline[i][0])

		if (self.cell.room.vector):
			room = self.cell.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(pline)):
				pline[i] = rotate(pline[i], rot, uv)

		return pline


	def draw_whole(self, msp):

		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = ay + self.height

		if (self.side==0 or self.side==2):
			if (self.mode==0):
				orig1 = (ax, by)
				orig2 = (bx, by)
			else:
				orig1 = (by, ax)
				orig2 = (by, bx)
		else:
			if (self.mode==0):
				orig1 = (ax, ay)
				orig2 = (bx, ay)
			else:
				orig1 = (ay, ax)
				orig2 = (ay, bx)

		if (self.side==0 or self.side==2):
			if (self.mode==0):
				rot1, xs1, ys1 = 90, 0.1, 0.1
				rot2, xs2, ys2 = 90, 0.1, -0.1
			else:
				rot1, xs1, ys1 = 0, 0.1, -0.1
				rot2, xs2, ys2 = 0, 0.1, 0.1
		else:
			if (self.mode==0):
				rot1, xs1, ys1 = 90, -0.1, 0.1
				rot2, xs2, ys2 = 90, -0.1, -0.1
			else:
				rot1, xs1, ys1 = 0, -0.1, -0.1
				rot2, xs2, ys2 = 0, -0.1, 0.1

		if (self.cell.room.vector):
			room = self.cell.room
			uv = room.uvector
			uv = (uv[0], -uv[1])
			orig1 = rotate(orig1, room.rot_orig, uv)
			orig2 = rotate(orig2, room.rot_orig, uv)
			rot1 += room.rot_angle
			rot2 += room.rot_angle

		xs1, ys1 = xs1/scale, ys1/scale
		xs2, ys2 = xs2/scale, ys2/scale

		block1 = msp.add_blockref(self.panel_type, orig1, 
			dxfattribs={'xscale': xs1, 'yscale': ys1, 'rotation': rot1})

		block2 = msp.add_blockref(self.panel_type, orig2, 
			dxfattribs={'xscale': xs2, 'yscale': ys2, 'rotation': rot2})

		block1.dxf.layer = layer_panel
		block2.dxf.layer = layer_panel


	def draw_half(self, msp):

		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = ay + self.height
	
		if (self.side==0):
			orig = (ax, by)
			if (self.mode==0):
				rot, xs, ys = 90, 0.1, 0.1
			else:
				rot, xs, ys = 0, 0.1, -0.1

		if (self.side==1):
			orig = (ax, ay)
			if (self.mode==0):
				rot, xs, ys = 90, -0.1, 0.1
			else:
				rot, xs, ys = 0, -0.1, -0.1

		if (self.side==2):
			orig = (bx,by)
			if (self.mode==0):
				rot, xs, ys = 90, 0.1, -0.1
			else:
				rot, xs, ys = 0, 0.1, 0.1

		if (self.side==3):
			orig = (bx,ay)
			if (self.mode==0):
				rot, xs, ys = 90, -0.1, -0.1
			else:
				rot, xs, ys = 0, -0.1, 0.1

		if (self.mode==1):
			orig = (orig[1], orig[0])

		if (self.cell.room.vector):
			room = self.cell.room
			uv = room.uvector
			uv = (uv[0], -uv[1])
			orig = rotate(orig, room.rot_orig, uv)
			rot += room.rot_angle

		xs, ys = xs/scale, ys/scale

		block = msp.add_blockref(self.panel_type, orig, 
			dxfattribs={'xscale': xs, 'yscale': ys, 'rotation': rot})

		block.dxf.layer = layer_panel

	def poly(self):

		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = self.ycoord + self.height

		pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]
		
		if (self.mode==1):
			for i in range(0,len(pline)):
				pline[i] = (pline[i][1], pline[i][0])

		if (self.cell.room.vector):
			room = self.cell.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(pline)):
				pline[i] = rotate(pline[i], rot, uv)

		return pline

	def lux_poly(self):

		global scale

		if not self.size[1] == 2:
			return

		if self.size == (2,2):
			ax = self.xcoord + 27.5/scale; bx = ax + 145/scale

		if self.size == (1,2) and (self.side==0 or self.side==1):
			ax = self.xcoord + 27.5/scale; bx = ax + 72.5/scale
			
		if self.size == (1,2) and (self.side==2 or self.side==3):
			ax = self.xcoord; bx = ax + 72.5/scale
		
		ay = self.ycoord + 51/scale; by = ay + 18/scale

		pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]

		if (self.mode==1):
			for i in range(0,len(pline)):
				pline[i] = (pline[i][1], pline[i][0])

		if (self.cell.room.vector):
			room = self.cell.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(pline)):
				pline[i] = rotate(pline[i], rot, uv)

		return pline


	def draw(self, msp):

		if (self.size[1]==2):
			if (self.color==bathroom_color):
				self.panel_type = block_green_120x100	
			else:
				self.panel_type = block_blue_120x100	
		else:
			if (self.color==bathroom_color):
				self.panel_type = block_green_60x100	
			else:
				self.panel_type = block_blue_60x100	

		if (self.size[0]==2):
			self.draw_whole(msp)
		else:
			self.draw_half(msp)

		if (self.stripe):
			self.draw_stripe(msp)


	def draw_stripe(self, msp):

		w = stripe_width/self.size[1] * self.height
		offs = stripe_offset/self.size[1]
		if (not self.side % 2):
			offs = 1-offs
		soffs = self.height*offs

		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord + soffs + w
		by = self.ycoord + soffs - w

		pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]
		
		if (self.mode==1):
			for i in range(0,len(pline)):
				pline[i] = (pline[i][1], pline[i][0])

		if (self.cell.room.vector):
			room = self.cell.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(pline)):
				pline[i] = rotate(pline[i], rot, uv)

		pl = msp.add_lwpolyline(pline)
		pl.dxf.layer = layer_panel
		pl.dxf.color = stripe_color
		
		hatch = msp.add_hatch(color=stripe_color)
		hatch.paths.add_polyline_path(pline, is_closed=True)
		hatch.dxf.layer = layer_panel

	def draw_profile(self, msp):
		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = ay + self.height
		dx = hatch_width; dy = hatch_height

		pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]
		
		if (self.size==(2,2) or self.size==(2,1)):
			if (self.side==1 or self.side==3):
				pline = [(ax,ay+dy),(ax,by),(bx,by),(bx,ay+dy),
					(bx-dx,ay+dy),(bx-dx,ay),(ax+dx,ay),(ax+dx,ay+dy),(ax,ay+dy)]

			if (self.side==0 or self.side==2):
				pline = [(ax,ay),(ax,by-dy),(ax+dx,by-dy),(ax+dx,by),(bx-dx,by),
					(bx-dx,by-dy),(bx,by-dy),(bx,ay),(ax,ay)]

		if (self.size==(1,1) or self.size==(1,2)):
			if (self.side==1):
				pline = [(ax,ay+dy),(ax,by),(bx,by),(bx,ay),(ax+dx,ay),
						 (ax+dx,ay+dy),(ax,ay+dy)]
			if (self.side==0):
				pline = [(ax,by-dy),(ax+dx,by-dy),(ax+dx,by),(bx,by),
						 (bx,ay),(ax,ay),(ax,by-dy)]
			if (self.side==3):
				pline = [(ax,ay),(ax,by),(bx,by),(bx,ay+dy),
						 (bx-dx,ay+dy),(bx-dx,ay),(ax,ay)]
			if (self.side==2):
				pline = [(ax,ay),(ax,by),(bx-dx,by),(bx-dx,by-dy),
						 (bx,by-dy),(bx,ay),(ax,ay)]


		if (self.mode==1):
			for i in range(0,len(pline)):
				pline[i] = (pline[i][1], pline[i][0])

		if (self.cell.room.vector):
			room = self.cell.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(pline)):
				pline[i] = rotate(pline[i], rot, uv)

		pl = msp.add_lwpolyline(pline)
		pl.dxf.layer = layer_panelp
		pl.dxf.color = self.color


	# Check if the distance from the wall is smaller 
	# than the minimum allowed and updates gapl and gapr
	def dist_from_walls(self):

		room = self.cell.room

		ax = self.xcoord; bx = ax + self.width
		ay = self.ycoord; by = ay + self.height

		if (self.side==0 or self.side==2):
			lft = (ax,by)
			rgt = (bx,by)
			self.gapl = room.hdist_from_poly(lft)
			self.gapr = room.hdist_from_poly(rgt)
		else:
			lft = (ax,ay)
			rgt = (bx,ay)
			self.gapl = room.hdist_from_poly(lft)
			self.gapr = room.hdist_from_poly(rgt)

		# dist from internal walls
		for obs in room.obstacles:
			self.gapl = min(obs.hdist_from_poly(lft), self.gapl)
			self.gapr = min(obs.hdist_from_poly(rgt), self.gapr)

class Dorsal:
	def __init__(self, grid, pos, side, room):
		self.panels = list()
		self.pos = pos
		self.grid = grid
		self.side = side
		self.room = room
		self.elems = 0
		self.gapl = MAX_DIST
		self.gapr = MAX_DIST
		self.gap = 0
		self.coupled = True

		m = self.grid
		self.cost = 0
		self.lost = 0
		j = self.pos[1] - 2
		i = self.pos[0]

		side = self.side
		self.first_of_line = True
		self.last_panel_cut = False
		self.last_handside = 0

		while(j<m.shape[1]-3):

			j += 2

			if (m[i,j] and m[i+1,j] and m[i,j+1] and m[i+1,j+1]):
				self.new_panel(m[i,j],(2,2),side)
				continue
		
			if side==1:
				if (m[i,j] and m[i,j+1] and
					(not m[i+1,j]) and (not m[i+1,j+1])):
					self.new_panel(m[i,j],(2,1))
					self.cost += 0.1
					continue

				# half panels vertical
				if m[i+1,j]:
					if m[i,j]:
						self.new_panel(m[i,j],(1,2),0)
						self.cost += 1
					else:
						self.cost += 4
						self.lost += 1	
				else:
					if m[i,j]:
						self.new_panel(m[i,j],(1,1),0)
						self.cost += 1

				if m[i+1,j+1]:
					if m[i,j+1]:
						self.new_panel(m[i,j+1],(1,2),1)
						self.cost += 1
					else:
						self.cost += 4
						self.lost += 1	
				else:
					if m[i,j+1]:
						self.new_panel(m[i,j+1],(1,1),1)
						self.cost += 1

			else:
				if (m[i+1,j] and m[i+1,j+1] and
					(not m[i,j]) and (not m[i,j+1])):
					self.new_panel(m[i+1,j],(2,1))
					self.cost += 0.1
					continue

				# half panels vertical
				if m[i,j]:
					if m[i+1,j]:
						self.new_panel(m[i,j],(1,2),0)
						self.cost += 1
					else:
						self.cost += 4
						self.lost += 1	
				else:
					if m[i+1,j]:
						self.new_panel(m[i+1,j],(1,1),0)
						self.cost += 1

				if m[i,j+1]:
					if m[i+1,j+1]:
						self.new_panel(m[i,j+1],(1,2),1)
						self.cost += 1
					else:
						self.cost += 4
						self.lost += 1	
				else:
					if m[i+1,j+1]:
						self.new_panel(m[i+1,j+1],(1,1),1)
						self.cost += 1

		# Now assign costs based on gaps
		if (self.last_panel_cut and self.room.clt_xside == RIGHT):
			self.cost += 0.01
			if (self.last_handside==0):
				self.cost += 0.005


	# handside 0=left, 1=right
	def new_panel(self, cell, size, handside=0):

		panel = Panel(cell,size,self.side+handside*2)
		self.elems += size[0]*size[1]

		panel.dist_from_walls()

		# Calculate water entry point
		ax = panel.xcoord; bx = ax + panel.width
		ay = panel.ycoord; by = ay + panel.height
		if (self.side==0):
			lft = (ax,by)
			rgt = (bx,by)
		else:
			lft = (ax,ay)
			rgt = (bx,ay)
		
		if (self.room.clt_xside == LEFT):
			if (len(self.panels)==0 or
			    self.attach[0] > lft[0]):
				self.attach = lft
		else:
			if (len(self.panels)==0 or
			    self.attach[0] < rgt[0]):
				self.attach = rgt

		if size==(2,2) or size==(2,1):
			self.last_panel_cut = False
		else:
			self.last_panel_cut = True
		self.last_hanside = handside

		if (self.first_of_line and self.last_panel_cut and
			  self.room.clt_xside==LEFT):
			self.cost += 0.01
			if (handside==1):
				self.cost += 0.005

		self.panels.append(panel)
		self.first_of_line = False

		if (panel.gapl < self.gapl):
			self.gapl = panel.gapl

		if (panel.gapr < self.gapr):
			self.gapr = panel.gapr



class Dorsals(list):
	def __init__(self, room):
		self.cost = 0 
		self.lost = 0
		self.gapl = MAX_DIST
		self.gapr = MAX_DIST
		self.gap = 0
		self.elems = 0
		self.panels = list()
		self.room = room
	
	def add(self, dorsal):	
		global min_dist
		self.cost += dorsal.cost
		self.elems += dorsal.elems
		self.panels += dorsal.panels

		if (len(dorsal.panels)>0):
			self.append(dorsal)
		
		self.gapr = min(self.gapr, dorsal.gapr)
		self.gapl = min(self.gapl, dorsal.gapl)

		if (self.room.clt_xside==LEFT):
			if (self.gapl<min_dist):
				return False
			else:
				self.gap = self.gapl

		if (self.room.clt_xside==RIGHT):
			if (self.gapr<min_dist):
				return False
			else:
				self.gap = self.gapr

		return True

	def __lt__(self, dorsals):
		if (self.lost<dorsals.lost):
			return True
		if (self.cost<dorsals.cost):
			return True
		return False

# Cell of the grid over which panels are laid out
class Cell:
	def __init__(self, pos, box, room, mode):
		self.pos = pos
		self.box = box
		self.room = room
		self.color = room.color
		self.mode = mode

	def draw(self, msp):
		box = self.box
		ax = box[0]; bx = box[1]
		ay = box[2]; by = box[3]

		pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]
		if (self.mode==1):
			for i in range(0,len(pline)):
				pline[i] = (pline[i][1], pline[i][0])

		if (self.room.vector):
			room = self.room
			rot = room.rot_orig
			uv = room.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(pline)):
				pline[i] = rotate(pline[i], rot, uv)

		pl = msp.add_lwpolyline(pline)
		pl.dxf.layer = layer_box


# This class represents the grid over which the panels
# are laid out.
class PanelArrangement:

	def __init__(self, room):
		self.best_grid = self.cells = list()
		self.dorsals = Dorsals(room)
		self.dorsals.cost = MAX_COST
		self.room = room
		self.elems = 0
		self.alloc_mode = -1
		self.circuits = list()
		self.strip_len = 0

	def len(self):
		return len(self.cells)

	def make_grid(self, origin):

		self.cells = list()

		(sx, sy) = origin
	
		col = 0
		sax = sx
		while(sax+panel_width <= self.room.bx):
			say = sy
			row = 0
			while(say+panel_height <= self.room.by):
				pos = (row, col)
				box = (sax, sax+panel_width, say, say+panel_height)
				flag = True
				if self.room.is_box_inside(box):
					for obstacle in self.room.obstacles:
						if not obstacle.is_box_outside(box):
							flag = False
							break
				else:
					flag = False

				if (flag):
					self.cells.append(Cell(pos, box, self.room, self.mode))
				say += panel_height
				row += 1

			sax += panel_width
			col += 1

		if (len(self.cells)==0):
			return False

		maxr = self.rows = max([c.pos[0] for c in self.cells]) + 7
		maxc = self.cols = max([c.pos[1] for c in self.cells]) + 3

		m = self.grid = np.full((maxr, maxc), None)

		for cell in self.cells:
			m[(cell.pos[0]+3, cell.pos[1]+1)] = cell


		return True

	def build_dorsals(self, pos):

		dorsals = Dorsals(self.room)

		i = pos[0]
		while(i<self.rows-3):

			bu0 = Dorsal(self.grid, (i,0), 0, self.room)
			bu1 = Dorsal(self.grid, (i,1), 0, self.room)
			td0 = Dorsal(self.grid, (i+2,0), 1, self.room)
			td1 = Dorsal(self.grid, (i+2,1), 1, self.room)

			b_dors = bu0
			t_dors = td0
			lost = b_dors.lost + t_dors.lost
			cost = b_dors.cost + t_dors.cost

			nlost = bu1.lost + td0.lost
			ncost = bu1.cost + td1.cost 


			if (nlost < lost or (nlost==lost and ncost<cost)):
				b_dors = bu1
				t_dors = td1
				cost = ncost
				lost = nlost

			if (b_dors.elems == 0 or t_dors.elems==0):
				b_dors.coupled = t_dors.coupled = False


			if (lost>0):

				bd0 = Dorsal(self.grid, (i,0), 1, self.room)
				bd1 = Dorsal(self.grid, (i,1), 1, self.room)
				tu0 = Dorsal(self.grid, (i+2,0), 0, self.room)
				tu1 = Dorsal(self.grid, (i+2,1), 0, self.room)

				bd = sorted([bu0, bu1, bd0, bd1], key=lambda x: (x.lost,x.cost))
				td = sorted([tu0, tu1, td0, td1], key=lambda x: (x.lost,x.cost))


				nlost = bd[0].lost + td[0].lost
				ncost = bd[0].cost + td[0].cost 
			
				if (nlost < lost):
					b_dors = bd[0]
					t_dors = td[0]
					b_dors.coupled = t_dors.coupled = False
					b_dors.cost += 0.5

			if (not dorsals.add(b_dors) or not dorsals.add(t_dors)):
				return None

			i += 4

		return dorsals


	
	def alloc_panels(self, origin):

		if (not self.make_grid(origin)):
			return

		#if (self.room.pindex==3):
		#	print(len(self.cells), self.room.clt_xside, self.room.clt_yside)
	
		#	print("room", self.room.pindex, self.mode, end="")
		#	if (self.room.clt_xside==LEFT):
		#		print(" LEFT ",end="")
		#	else:
		#		print(" RIGHT ",end="")
		#	if (self.room.clt_yside==TOP):
		#		print("TOP")
		#	else:
		#		print("BOTTOM")

		#	for row in range(self.rows):
		#		for col in range(self.cols):
		#			if (self.grid[row,col]):
		#				print("X", end="")
		#			else:
		#				print(".", end="")
		#		print()

		#for j in range(0,2):
		for i in range(0,4):
			# set origin of dorsals
			pos = (i, 0)
			trial_dorsals = self.build_dorsals(pos)
			if ((not trial_dorsals == None ) and
				((trial_dorsals.elems > self.dorsals.elems) or
				(trial_dorsals.elems == self.dorsals.elems and
				 trial_dorsals.cost < self.dorsals.cost) or
				(trial_dorsals.elems == self.dorsals.elems and
				 trial_dorsals.cost == self.dorsals.cost and
				 trial_dorsals.gap > self.dorsals.gap))):
					self.dorsals = trial_dorsals
					self.best_grid = self.cells
					self.alloc_mode = self.mode
					self.alloc_clt_xside = self.room.clt_xside
					self.origin = origin


	def fittings(self):

		cside = self.alloc_clt_xside
		fits = self.fittings = list()

		for dorsal in self.dorsals:

			local_fit = list()
			for p in dorsal.panels:

				y, x = p.cell.pos
				w = p.size[0]; h = p.size[1]
				btm = (1 - p.side % 2)
				lft = p.side//2

				cpos = x + w*lft, y + h*btm
				local_fit.append([{'pos': cpos,
							'side': lft,
							'hside': btm,
							'dorsal': dorsal,
							'panel': p,
					  		'last': False }])

				if (p.size[0]==2):
					cpos = x + w*(1-lft), y + h*btm
					local_fit.append([{'pos': cpos,
								  'side': 1-lft,
								  'hside': btm,
								  'dorsal': dorsal,
								  'panel': p,
							      'last': False }])

			local_fit.sort(key=lambda x: x[0]['pos'])

			if (cside==0):
				local_fit[-1][0]["last"] = True
			else:
				local_fit[0][0]["last"] = True

			fits += local_fit

		# merge fittings 
		for j, fit in enumerate(fits):
			fit[0]['merged'] = False
			for k in range(j):
				fitv = fits[k]
				if (not fitv[0]['merged'] and 
					fitv[0]['pos'] == fit[0]['pos']):
					fitv.append(fit[0])
					fit[0]['merged'] = True

		self._fits = list(filter(lambda x: not x[0]['merged'], fits))


	def make_couplings(self):

		# create coupling from merged fittings
		cpls = self.couplings = list()
		for fit in self._fits:
			cpls.append(Coupling(fit))

		# form circuits from couplings
		# and split couplings
		cirs = self.circuits = []
		ypos = -1
		for cp in cpls:
			if (cp.ypos > ypos):
				ypos = cp.ypos
				cirs.append(Circuit(self.room))
			cirs[-1].add_coupling(cp)


		self.couplings = []
		self.strip_len = 0
		for cir in cirs:

			# calculate couplings
			self.couplings += cir.name_couplings()

			# add red stripe
			if (not cir.is_double()):
				self.strip_len += cir.xmax - cir.xmin
				cir.add_stripe()

		#print("circuits, room", self.room.pindex)
		#for cir in self.circuits:
		#	print("[",cir.xmin, cir.xmax,"]")
		#	for cpl in cir.couplings:
		#		print(cpl.type)

		#print("strip_len", self.strip_len)
		#print()


	def draw_grid(self, msp):	
		for cell in self.best_grid:
			cell.draw(msp)



class Coupling:
	def __init__(self, fits):
		self.fits = fits
		self.xpos = fits[0]['pos'][0]
		self.ypos = fits[0]['pos'][1]
		self.pos = fits[0]['pos']
		self.num_fits = len(fits)
		self.flip = False

		self.is_double = False
		first_fit = fits[0]['dorsal']	
		for fit in fits:
			if (fit['dorsal'] != first_fit):
				self.is_double = True

		top_last = bottom_last = None
		for fit in fits:
			if (fit['hside']==0):
				if top_last == None:
					top_last = fit['last']
				else:
					top_last |= fit['last']

			if (fit['hside']==1):
				if bottom_last == None:
					bottom_last = fit['last']
				else:
					bottom_last |= fit['last']

		if top_last == None:
			top_last = True

		if bottom_last == None:
			bottom_last = True
 
		self.is_last = top_last and bottom_last
		#if (self.pos == (6,6)):
		#	print("STRANGE CASE: ", top_last, bottom_last, self.is_last)


	def add_stripe(self, xmin, xmax):
		for fit in self.fits:
			panel = fit['panel']
			xa = panel.cell.pos[1]
			xb = xa + panel.size[0]
			if (not (xa<xmin or xb>xmax)):
				fit['panel'].stripe = True

	def print(self):
		for fit in self.fits:
			print(fit['pos'], end=" ")
		print()

	def is_at_top(self):
		return self.fits[0]['hside']

	def is_at_right(self):
		return self.fits[0]['side']

class Circuit:
	def __init__(self, room):
		self.room = room
		self.couplings = list()
		self.panels = list()
		self.size = 0
		self.flip = False
		self.xa = self.xb = None

	def add_coupling(self, cpl):
		self.couplings.append(cpl)
		cside = self.room.arrangement.alloc_clt_xside

		for fit in cpl.fits:
			panel = fit['panel']
			if not panel in self.panels:
				self.panels.append(panel)
				size = panel.size
				self.size += (size[0]*size[1])/4

				# calculating front line
				xa = panel.cell.pos[1]
				xb = xa + size[0]
				ya = panel.cell.pos[0]
				yb = ya + size[1]
				if (self.xa == None):
					self.xa = xa
					self.xb = xb
					self.ya = ya
					self.yb = yb
					self.xa_panel = panel
					self.xb_panel = panel
				else:
					if xa == self.xa and cside == LEFT:
						if (yb > self.yb):
							self.yb = yb
						if (ya < self.ya):
							self.ya = ya
						
					if xb == self.xb and cside == RIGHT:
						if (yb > self.yb):
							self.yb = yb
						if (ya < self.ya):
							self.ya = ya

					if xa < self.xa:
						self.xa = xa
						if cside == LEFT:
							self.ya = ya
							self.yb = yb
						self.xa_panel = panel

					if xb > self.xb:
						self.xb = xb
						if cside == RIGHT:
							self.ya = ya
							self.yb = yb
						self.xb_panel = panel 

	def is_double(self):

		for cpl in self.couplings:
			if (cpl.is_double):
				return True
		return False


	def add_stripe(self):
		for cpl in self.couplings:
			cpl.add_stripe(self.xmin, self.xmax)


	def face(self):
		if self.flip:
			return [(self.xa, self.ya), (self.xa, self.yb)]
		else:
		 	return [(self.xb, self.ya), (self.xb, self.yb)]

	def name_couplings(self):	

		cside = self.room.arrangement.alloc_clt_xside
		if cside == LEFT:
			self.flip = True

		cpls = self.couplings	

		self.xmin = self.xmax = cpls[0].xpos
		self.ypos = cpls[0].ypos
		cplmin = cplmax = cpls[0]
		for cpl in cpls:
			if (cpl.xpos < self.xmin):
				self.xmin = cpl.xpos
				cplmin = cpl

			if (cpl.xpos > self.xmax):
				self.xmax = cpl.xpos
				cplmax = cpl

		self.fixture = None
		if (cside==RIGHT and self.xb!=cplmax.xpos):
			self.fixture = self.xb
			pos = self.xb, self.ypos
			fit = {'pos' : pos, 
				   'dorsal': None,
				   'hside': 0,
				   'side': 0,
				   'last': 0,
				   'panel': self.xb_panel
			}
			cpl = Coupling([fit])
			cpl.type = 'Rac_20_20_dritto'
			cpls.append(cpl)

		if (cside==LEFT and self.xa!=cplmin.xpos):
			self.fixture = self.xa
			pos = self.xa, self.ypos
			fit = {'pos' : pos, 
				   'dorsal': None,
				   'hside': 0,
				   'side': 0,
				   'last': 0,
				   'panel': self.xa_panel
			}
			cpl = Coupling([fit])
			cpl.type = 'Rac_20_20_dritto'
			cpls.append(cpl)

		for cpl in cpls:
			cpl.circuit = self

			if hasattr(cpl,"type"):
				continue

			end_flag = False
			if (cpl==cplmin or cpl==cplmax) and cpl.is_last==True:
				end_flag = True
				cpl.flip = self.flip

			if (cpl.num_fits==1): 
				if (end_flag):
					cpl.type = "Rac_20_10"
				else:
					cpl.type = "Rac_20_10_20"

			if (cpl.num_fits==2):

				if (cpl.is_double):
					f0 = cpl.fits[0]
					f1 = cpl.fits[1]
					if f0['side'] != f1['side']:
						# Double fitting on opposite sides. Split
						cpl.type = "invalid"
						cp0 = Coupling([f0])
						cp1 = Coupling([f1])
						cpls.append(cp0)
						cpls.append(cp1)
					else:
						# Double fitting on the same side.
						if (end_flag):
							cpl.type = "Rac_10_20_10"
						else:
							cpl.type = "Rac_20_10_20_10"
				else:
					if (end_flag):
						cpl.type = "Rac_20_10_10"
					else:
						cpl.type = "Rac_20_10_10_20"

			if (cpl.num_fits==3):

				if (end_flag):
					cpl.type = "Rac_10_20_10_10"
				else:
					# Split if not an ending
					cpl.type = "invalid"
					
					f0 = cpl.fits[0]
					f1 = cpl.fits[1]
					f2 = cpl.fits[2]
					if f0['side'] == f1['side']:
						cp0 = Coupling([f0, f1])
						cp1 = Coupling([f2])
					else:
						if f0['side'] == f1['side']:
							cp0 = Coupling([f0, f2])
							cp1 = Coupling([f1])
						else:
							cp0 = Coupling([f1, f2])
							cp1 = Coupling([f0])
					cpls.append(cp0)
					cpls.append(cp1)

			if (cpl.num_fits==4):
				if (end_flag):
					cpl.type = "Rac_10_10_20_10_10"
				else:
					cpl.type = "Rac_20_10_10_20_10_10"
	
		cpls.sort(reverse=(cside==LEFT), key=lambda x: x.xpos)

		return cpls

				
	def print(self):

		for cp in self.couplings:
			print(cp.type, end="")
			print(", ", end="")
		print()


class Line:
	def __init__(self, room, joined=False):
		self.couplings = list()
		self.flow = 0
		self.room = room
		self.sgn = 1
		self.facing = False
		
		# if a line is joined then 
		# contains a list of lines
		self.joined = joined
		self.lines = list()

	def draw_couplings(self, msp):
		global panel_width, panel_height
	
		if (not hasattr(self, 'couplings')):
			return	

		room = self.room
		arrng = room.arrangement

		w = panel_width
		h = panel_height
		x0, y0 = arrng.origin
		for cpl in self.couplings:

			if cpl.type == "invalid":
				continue

			#print(cpl.type, cpl.pos, end="")
			#print("Room=", self.room.pindex, end="")
			#print(" mode=", self.alloc_mode, end="")
			#print(" vect=", self.room.vector)

			xpos, ypos = cpl.pos
			fit = fittings[cpl.type]

			symbol = deepcopy(fit["symbol"])

			sgnx = -1; sgny = -1
			# flipping symbol upside down
			if ((cpl.type == "Rac_20_10_10_20" or 
				cpl.type == "Rac_20_10_10"  or
				cpl.type == "Rac_20_10_20"  or
				cpl.type == "Rac_20_10") and
				cpl.is_at_top()):
				sgny = -sgny
				#print("flipping upside-down")
				for i, pline in enumerate(symbol):
					for k, p in enumerate(pline):
						x, y = p
						symbol[i][k] = x, -y

			# flipping symbol left-right 
			if ((cpl.type == "Rac_10_10_20_10_10" or
				cpl.type == "Rac_10_20_10"  or
				cpl.type == "Rac_20_10_10" or
				cpl.type == "Rac_20_10") and
				cpl.flip):
				sgnx = -sgnx

				for i, pline in enumerate(symbol):
					for k, p in enumerate(pline):
						x, y = p
						symbol[i][k] = -x, y
				
			# horizontal offset
			offset = 0
			if (cpl.type == "Rac_10_20_10"  or
				cpl.type == "Rac_20_10_20_10" or
				cpl.type == "Rac_20_10_20" or
				cpl.type == "Rac_20_10"):
				offset =  shift * (0.5 - cpl.is_at_right())
				for i, pline in enumerate(symbol):
					for k, p in enumerate(pline):
						x, y = p
						symbol[i][k] = x + offset, y	

			# shift axis if single-sided dorsal
			axis = 0
			if not cpl.circuit.is_double():
				axis = axis_offset*(0.5-cpl.is_at_top())


			xo   = x0 + w*xpos + offset/scale
			yo_1 = y0 + h*ypos + (axis+delta_v)/scale
			yo_2 = y0 + h*ypos + (axis-delta_v)/scale

			if (arrng.alloc_mode == 0):
				orig1 = xo, yo_1
				orig2 = xo, yo_2
				rot = 0
			else:
				orig1 = yo_1, xo
				orig2 = yo_2, xo
				sgny = -sgny
				rot = 90

			if (self.room.vector):
				room = self.room
				rot_orig = room.rot_orig
				uv = room.uvector
				uv = (uv[0], -uv[1])
				orig1 = rotate(orig1, rot_orig, uv)
				orig2 = rotate(orig2, rot_orig, uv)
				rot += room.rot_angle

			xs, ys = sgnx*0.1/scale, sgny*0.1/scale
			if (cpl.type != 'Rac_20_20_dritto'):
				blk_red = msp.add_blockref(cpl.type + "_rosso", orig1,
					dxfattribs={'xscale': xs, 'yscale': ys, 'rotation': rot})
				blk_blu = msp.add_blockref(cpl.type + "_blu", orig2,
					dxfattribs={'xscale': xs, 'yscale': ys, 'rotation': rot})
			else:
				bkl_red = msp.add_blockref(cpl.type, orig1,
					dxfattribs={'xscale': xs, 'yscale': ys, 'rotation': rot})
				blk_blu = msp.add_blockref(cpl.type, orig2,
					dxfattribs={'xscale': xs, 'yscale': ys, 'rotation': rot})
			
			blk_red.dxf.layer = layer_joints
			blk_blu.dxf.layer = layer_joints
			cpl.orig1 = orig1
			cpl.orig2 = orig2

			if (not debug):
				continue

			# draw lines
			for pline in symbol:
				spline1 = []
				spline2 = []
				for p in pline:	
					xp1 = x0 + (p[0]+delta_h)/scale + w*xpos
					yp1 = y0 + (p[1]+delta_v+axis)/scale + h*ypos
					xp2 = x0 + (p[0]-delta_h)/scale + w*xpos
					yp2 = y0 + (p[1]-delta_v+axis)/scale + h*ypos
					if (arrng.alloc_mode == 0):
						ps1 = xp1, yp1
						ps2 = xp2, yp2
					else:
						ps1 = yp1, xp1
						ps2 = yp2, xp2
					spline1.append(ps1)
					spline2.append(ps2)
				
				if (room.vector):
					rot = room.rot_orig
					uv = room.uvector
					uv = (uv[0], -uv[1])
					for i in range(len(spline1)):
						spline1[i] = rotate(spline1[i], rot, uv)
						spline2[i] = rotate(spline2[i], rot, uv)

				pl = msp.add_lwpolyline(spline1)
				pl.dxf.layer = layer_joints
				pl.dxf.color = color_warm
				pl.dxf.lineweight = 2

				pl = msp.add_lwpolyline(spline2)
				pl.dxf.layer = layer_joints
				pl.dxf.color = color_cold
				pl.dxf.lineweight = 2

	def draw_cpls_link(self, msp, cpls):

		for i in range(len(cpls)-1):
			seg1 = [cpls[i].orig1, cpls[i+1].orig1]
			pl1 = msp.add_lwpolyline(seg1)
			pl1.dxf.color = color_warm
			pl1.dxf.layer = layer_joints

			seg2 = [cpls[i].orig2, cpls[i+1].orig2]
			pl2 = msp.add_lwpolyline(seg2)
			pl2.dxf.color = color_cold
			pl2.dxf.layer = layer_joints

	def draw_adductions(self, msp):

		if not self.joined:
			self.draw_cpls_link(msp, self.couplings)

		else:
			for line in self.lines:
				self.draw_cpls_link(msp, line.couplings)				


# This class represents the room described by a polyline
class Room:

	index = 1

	def __init__(self, poly, output):

		self.poly = poly
		self.index = Room.index
		Room.index = Room.index + 1
		self.output = output
		self.ignore = False
		self.errorstr = ""
		self.contained_in = None
		self.obstacles = list()
		self.gates = list()
		self.lines = list()
		self.joined_lines = list()
		self.sup = 0
		self.inf = 0
		self.fixed_collector = None

		tol = tolerance
		self.orient = 0
		self.boxes = list()
		self.coord = list()

		self.points = list(poly.vertices())	
		self.vector = None
		self.vector_auto = False

		# Add a final point to closed polylines
		p = self.points
		if (poly.is_closed):
			self.points.append((p[0][0], p[0][1]))

		# Check if the polyine is open with large final gap
		n = len(p)-1
		if (abs(p[0][0]-p[n][0])>tol or abs(p[0][1]-p[n][1])>tol):
			self.errorstr = "Error: open polyline in Zone %d \n" % self.index 
			self.ignore = True
			return

		# Straighten up walls
		finished = False
		while not finished:
			for i in range(1, len(p)):
				if (abs(p[i][0]-p[i-1][0]) < tol):
					p[i] = (p[i-1][0], p[i][1])

				if (abs(p[i][1]-p[i-1][1]) < tol):
					p[i] = (p[i][0], p[i-1][1])

			if (p[0] == p[n]):
				finished = True
			else:
				p[0] = p[n]


		self.color = poly.dxf.color
		self.arrangement = PanelArrangement(self)
		self.panels = list()
		self.bounding_box()
		self.area = self._area()
		self.pos = self._barycentre()
		self.perimeter = self._perimeter()
		self.collector = None

	def _area(self):
		a = 0
		p = self.points
		for i in range(0, len(p)-1):
			a += (p[i+1][0]-p[i][0])*(p[i+1][1] + p[i][1])/2
		return abs(a/10000)

	def _centre(self):
		(cx, cy) = (0, 0)
		p = self.points
		n = len(p)-1
		for i in range(0,n):
			p = self.points[i]
			(cx, cy) = (cx+p[0], cy+p[1])

		return (cx/n, cy/n)
		

	def _perimeter(self):
		p = self.points
		d = 0
		for i in range(0, len(p)-1):
			d += sqrt(pow(p[i+1][0]-p[i][0],2)+pow(p[i+1][1]-p[i][1],2))
		return d

	def _barycentre(self):
		
		p = self.points
		(xb, yb) = (0, 0)
		Ax = Ay = 0
		for i in range(len(p)-1):
			(x0, y0) = p[i]
			(x1, y1) = p[i+1]
			xb += (y1-y0)*(x0*x0+x0*x1+x1*x1)/6
			yb += (x1-x0)*(y0*y0+y0*y1+y1*y1)/6
			Ax += (y1-y0)*(x0+x1)/2
			Ay += (x1-x0)*(y0+y1)/2
		
		if (Ax!=0 and Ay!=0):
			return (xb/Ax, yb/Ay)
		else:
			return (0, 0)

	def bounding_box(self):

		# Projections of coordinates on x and y
		self.xcoord = sorted(set([p[0] for p in self.points]))	
		self.ycoord = sorted(set([p[1] for p in self.points]))	

		self.ax = min(self.xcoord)
		self.bx = max(self.xcoord)
		self.ay = min(self.ycoord)
		self.by = max(self.ycoord)

	# Cell to absolute coordinates
	def abs(self, pos):
		global panel_width, panel_height
	
		arrng = self.arrangement

		w = panel_width
		h = panel_height
		x0, y0 = arrng.origin

		xo = x0 + w*pos[0]
		yo = y0 + h*pos[1]

		if (arrng.alloc_mode == 0):
			p = xo, yo
		else:
			p = yo, xo

		if (self.vector):
			rot_orig = self.rot_orig
			uv = self.uvector
			uv = (uv[0], -uv[1])
			p = rotate(p, rot_orig, uv)

		return p
		
	def dist_linear(self, collector):

		dorsals = self.arrangement.dorsals
		d = MAX_DIST
		for dorsal in dorsals:
			cltr_dist = dist(dorsal.pos, collector.pos)
			if (d > cltr_dist):
				d = cltr_dist 
				attachment = dorsal.pos
		return (d, attachment)

	def dist_on_tree(self, collector):

		dorsals = self.arrangement.dorsals
		d = MAX_DIST
		attachment = None
		for dorsal in dorsals:
			cltr_dist = dist(dorsal.pos, collector.pos)
			if (d >= cltr_dist):
				d = cltr_dist 
				attachment = dorsal.pos

		d = MAX_DIST
		uplink = None
		for cltr, walk, next_room in self.links:
			if (cltr == collector):
				d = walk
				uplink = next_room

		return (d, attachment, uplink)


	def wall(self, room):
		for nroom, nwall in self.gates:
			if (room==nroom):
				return nwall

	def orient_room(self):
		global max_room_area

		vtx = [(p[0],p[1],0) for p in self.points]
		conv_hull = convex_hull_2d(vtx)
		ch = [(s.x, s.y) for s in conv_hull]
		ch = [*ch, ch[0]]

		max_area = MAX_DIST2
		max_uv = (1,0)
		for i in range(len(ch)-1):
			p0, p1 = ch[i], ch[i+1]
			norm_uv = dist(p0, p1)
			if (norm_uv == 0):
				continue
			uvx, uvy = uv = (p1[0]-p0[0])/norm_uv, (p1[1]-p0[1])/norm_uv
			bxm = bxM = 0; by = 0
			for p in ch[:-1]:
				px =  uvx*(p[0]-p0[0]) + uvy*(p[1]-p0[1])
				py = abs(-uvy*(p[0]-p0[0]) + uvx*(p[1]-p0[1]))
				if (py > by):
					by = py

				if (px < bxm):
					bxm = px
				if (px > bxM):
					bxM = px

			Ar = (bxM - bxm)*by
			if ( Ar < max_area ):
				max_area = Ar
				max_uv = uv
				max_rot_orig = p0

		angle = min(abs(uvx),abs(uvy))/max(abs(uvx),abs(uvy))
		if (angle > 0.01):
			self.vector = True
			self.vector_auto = True
			self.uvector = max_uv
			self.rot_orig = max_rot_orig
			self.rot_angle = -atan2(max_uv[1], -max_uv[0])*180/pi
				

	# Building Room
	def alloc_panels(self):
		global panel_height, panel_width, search_tol

		# Rotate according to vector
		if (self.vector):

			p = self.points
			rot = self.rot_orig
			uv = self.uvector

			for i in range(len(p)):
				p[i] = rotate(p[i], rot, uv)

			for obs in self.obstacles:
				for i in range(len(obs.points)):
					obs.points[i] = rotate(obs.points[i], rot, uv)

			# collector position
			if (self.vector_auto):
				# program found vector
				cltr = self.collector
				rpos = rotate(self.pos, rot, uv)
				cpos = rotate(cltr.pos, rot, uv)

				(vx, vy) = (cpos[0]-rpos[0], cpos[1]-rpos[1])

				if (vx>=0):
					self.clt_xside = RIGHT
				else:
					self.clt_xside = LEFT

				if (vy>=0):
					self.clt_yside = TOP
				else:
					self.clt_yside = BOTTOM
			else:
				# user added vector
				self.clt_xside = LEFT 
				self.clt_yside = BOTTOM 

		self.arrangement.mode = 0   ;# horizontal
		while self.arrangement.mode < 2:
			self.bounding_box()
			# search within a small panel range
			sx = 0
			while sx<panel_width:
				sy = 0
				while sy<panel_height:
					origin = (sx + self.ax, sy + self.ay)
					self.arrangement.alloc_panels(origin)
					self.panels = self.arrangement.dorsals.panels
				
					sy += search_tol
				sx += search_tol

			if (self.vector and not self.vector_auto):
				break

			for i in range(0,len(self.points)):
				self.points[i] = (self.points[i][1], self.points[i][0])
			(self.xcoord, self.ycoord) = (self.ycoord, self.xcoord)

			for obs in self.obstacles:
				for i in range(0,len(obs.points)):
					obs.points[i] = (obs.points[i][1], obs.points[i][0])
				(obsxcoord, obs.ycoord) = (obs.ycoord, obs.xcoord)

			(self.clt_xside,self.clt_yside) = (self.clt_yside,self.clt_xside) 
	
			self.arrangement.mode += 1  

		# Rotate back according to vector
		if (self.vector):
			p = self.points
			uv = self.uvector
			uv = (uv[0], -uv[1])
			for i in range(len(p)):
				p[i] = rotate(p[i], rot, uv)

			for obs in self.obstacles:
				for i in range(len(obs.points)):
					obs.points[i] = rotate(obs.points[i], rot, uv)

		self.bounding_box()


	def is_point_inside(self, point):


		p = self.points
		x, y = point
		ints = 0

		for i in range(0, len(p)-1):
	
			if (p[i][0]==x and p[i+1][0]==x):
				if (max(p[i][1], p[i+1][1]) >= y
				  and min(p[i][1], p[i+1][1]) <=y):
					return True

			#if (p[i][1]==y and p[i+1][1]==y and
			#	min(p[i][0], p[i+1][0]) <= x <= max(p[i][0], p[i+1][0])):
			#		return True

			if (p[i][0] == p[i+1][0]):
				continue	

			if (p[i][0] < p[i+1][0]):
				x0, y0 = p[i][0], p[i][1]
				x1, y1 = p[i+1][0], p[i+1][1]
			else:
				x0, y0 = p[i+1][0], p[i+1][1]
				x1, y1 = p[i][0], p[i][1]

			
			if (x0<=x and x<x1):
				if (y0==y1 and y0==y):
					return True

				delta = x1 - x0
				py = (y0*(x1-x) - y1*(x0-x))/delta
				if (py>=y):
					ints += 1

		return (ints % 2 == 1)


	def is_box_inside(self, box):

		# Check if the center is inside
		center = ((box[0]+box[1])/2, (box[2]+box[3])/2)
		if (not self.is_point_inside(center)):
			return False

		# Does polyline cross box?
		p = self.points
		for i in range(0, len(p)-1):
			line = (p[i], p[i+1]) 
			if (cross(box, line)):
				return False

		return True

	def is_box_outside(self, box):

		# Check if the center is inside
		center = ((box[0]+box[1])/2, (box[2]+box[3])/2)
		if (self.is_point_inside(center)):
			return False

		# Does polyline cross box?
		p = self.points
		for i in range(0, len(p)-1):
			line = (p[i], p[i+1]) 
			if (cross(box, line)):
				return False

		return True

	def dist2_from_poly(self, point):

		cd = MAX_DIST
		p = self.points
		for i in range(0, len(p)-1):
			line = (p[i], p[i+1]) 
			d2 = dist2(line, point)
			if (d2<cd):
				cd = d2
		return cd

	def hdist_from_poly(self, point):

		cd = MAX_DIST
		p = self.points
		for i in range(0, len(p)-1):
			line = (p[i], p[i+1]) 
			d2 = hdist(line, point)
			if (d2<cd):
				cd = d2
		return cd

	def contains(self, room):
		for point in room.points:
			if (self.is_point_inside(point)):
				return True

		return False


	def collides_with(self, room):

		# check if BBs overlap
		if (self.ax > room.bx or
			self.bx < room.ax or
			self.ay > room.by or
			self.by < room.ay):
			return False

		# check if room contains self or is contained
		if (room.contains(self) or self.contains(room)):
			return True

		# check if polylines cross each other
		p = self.points
		q = room.points
		for i in range(0, len(p)-1):
			line1 = (Vec2(p[i]), Vec2(p[i+1]))
			for j in range(0, len(q)-1):
				line2 = (Vec2(q[j]), Vec2(q[j+1]))
				if (intersection_line_line_2d(line1, line2, virtual=False)):
					return True

		return False
		

	def add_gates(self, room):
		
		p1 = self.points
		p2 = room.points
		for i in range(0, len(p1)-1):
			line1 = (p1[i], p1[i+1]) 
			for j in range(0, len(p2)-1):
				line2 = (p2[j], p2[j+1]) 
				cond, wall = is_gate(line2, line1)
				if (cond):
					self.gates.append((room, wall))

	def set_as_root(self, queue, collector):
		self.visited = True

		for room, wall in self.gates:
			if (not room.visited):
				if (collector.contained_in == self):
					pos = collector.pos
				else:
					pos = self.pos
				distance = dist(pos, room.pos)
				walk = self.walk + distance
				if (walk < room.walk):
					room.walk = walk
					room.uplink = self

		# select next room
		if (len(queue)>0):
			queue.sort(key=lambda x: x.walk)
			next_room = queue.pop(0)
			next_room.set_as_root(queue, collector)

	def contains_vector(self, v):
		p1 = (v.dxf.start[0], v.dxf.start[1])
		p2 = (v.dxf.end[0], v.dxf.end[1])
		if (p1 == p2):
			return False
		return self.is_point_inside(p1) and self.is_point_inside(p2)


	def make_lines(self, fline, fpanel):

		lines_res = list()
		for cir in self.arrangement.circuits:
			cir.flow = cir.size * fpanel

			line = Line(self)
			line.sgn = 1
			line.level = cir.couplings[-1].pos[1]
			if cir.flip:
				line.sgn = -1
			self.lines.append(line)

			for cpl in cir.couplings:
				if (cpl.type == "invalid"):
					continue

				cpl.flow = 0	
				for fit in cpl.fits:
					fitflow = fpanel * fit["panel"].size[1]/4
					cpl.flow += fitflow

				if line.flow + cpl.flow > fline:
					if (cpl.type[-5:] == "_open"):
						cpl.type = cpl.type[:-5] + "_end"
						cpl.flip = cir.flip
					line = Line(self)
					line.sgn = 1
					line.level = cir.couplings[-1].pos[1]
					if cir.flip:
						line.sgn = -1
					self.lines.append(line)

				line.couplings.append(cpl)
				line.flow += cpl.flow

			line = self.lines.pop()
			line.facing = True
			lines_res.append(line)

		lines_res.sort(key=lambda x: x.flow)

		if len(lines_res)==0:
			return

		# join residual lines using first-fit decreasing
		joined_lines = list()
		while lines_res:
			line = lines_res.pop()
			for l in joined_lines:
				if l.flow + line.flow < fline:
					l.couplings += line.couplings.copy()
					l.flow += line.flow
					l.lines += [line]
					break
			else:
				nl = Line(self, joined=True)
				nl.lines.append(line)
				nl.couplings = line.couplings.copy()
				nl.flow = line.flow
				nl.level = line.couplings[-1].pos[1]
				joined_lines.append(nl)

	
		for line in joined_lines:
			if line.joined and len(line.lines)==1:
				line.lines[0].joined = False
				self.lines.append(line.lines[0])
			else:
				self.lines.append(line)


	# Reporting Room
	def report(self):
		txt = ""
		p2x2 = 0
		p2x1 = 0
		p1x2 = 0
		p1x1 = 0

		p1x1_r = 0
		p1x1_l = 0
		p1x2_r = 0
		p1x2_l = 0

		arrng = self.arrangement
		w = default_panel_width
		h = default_panel_height

		for panel in self.panels:
			if (panel.size == (2,2)):
				p2x2 += 1
			if (panel.size == (2,1)):
				p2x1 += 1
			if (panel.size == (1,2)):
				p1x2 += 1
				if (panel.side == 0 or panel.side == 3):
					if (arrng.alloc_mode == 0):
						p1x2_l += 1
					else:
						p1x2_r += 1	
				else:
					if (arrng.alloc_mode == 0):
						p1x2_r += 1
					else:
						p1x2_l += 1	

			if (panel.size == (1,1)):
				p1x1 += 1
				if (panel.side == 0 or panel.side == 3):
					if (arrng.alloc_mode == 0):
						p1x1_l += 1
					else:
						p1x1_r += 1
				else:
					if (arrng.alloc_mode == 0):
						p1x1_r += 1
					else:
						p1x1_l += 1

		area = self.area * scale * scale	
		active_area = w*h*(4*p2x2 + 2*p2x1 + 2*p1x2 + p1x1)/10000
		active_ratio = 100*active_area/area

		self.area_m2 = area
		self.active_m2 = active_area
		self.ratio = active_ratio

		txt += "Room area: %.4g m2 \n" % area
		txt += "Active area: %.4g m2 (%.4g%%)\n" % (active_area, active_ratio)
		txt += "  %5d panels %dx%d cm\n" % (p2x2, 2*w, 2*h) 
		txt += "  %5d panels %dx%d cm\n" % (p2x1, 2*w, h) 
		txt += "  %5d panels %dx%d cm - " % (p1x2, w, 2*h) 
		txt += " %d left, %d right\n" % (p1x2_l, p1x2_r)
		txt += "  %5d panels %dx%d cm - " % (p1x1, w, h) 
		txt += " %d left, %d right\n" % (p1x1_l, p1x1_r)


		self.room_rep = {
			"txt":               txt,
			"area":              area,
			"active_area":       active_area,
			"panels_120x200":    p2x2,
			"panels_60x200":     p2x1,
			"panels_120x100":    p1x2,
			"panels_120x100_l":  p1x2_l,
			"panels_120x100_r":  p1x2_r,
			"panels_60x100":     p1x1,
			"panels_60x100_l":   p1x1_l,
			"panels_60x100_r":   p1x1_r,
		}

		return self.room_rep


	def recount_panels(self, msp):

		txt = ""
		p2x2 = 0
		p2x1 = 0
		p1x2 = 0
		p1x1 = 0

		p1x1_r = 0
		p1x1_l = 0
		p1x2_r = 0
		p1x2_l = 0

		arrng = self.arrangement
		w = default_panel_width
		h = default_panel_height

		for e in msp.query('*[layer=="%s"]' % layer_panel):
			if (e.dxftype() == 'INSERT'):
				loc = e.dxf.insert
				location = loc[0], loc[1]
				if self.is_point_inside(location):
					vals = e.block().name.split("_")
					if vals[0] == "LEO":
						sgn = (e.dxf.xscale * e.dxf.yscale > 0)
						if vals[2] == "120":
							p1x2 += 1
							if sgn:
								p1x2_l += 1
							else:
								p1x2_r += 1

						if vals[2] == "60":
							p1x1 += 1
							if sgn:
								p1x1_l += 1
							else:
								p1x1_r += 1

		area = self.area * scale * scale	
		active_area = w*h*(4*p2x2 + 2*p2x1 + 2*p1x2 + p1x1)/10000
		active_ratio = 100*active_area/area

		self.area_m2 = area
		self.active_m2 = active_area
		self.ratio = active_ratio

		txt += "Room area: %.4g m2 \n" % area
		txt += "Active area: %.4g m2 (%.4g%%)\n" % (active_area, active_ratio)
		txt += "  %5d panels %dx%d cm\n" % (p2x2, 2*w, 2*h) 
		txt += "  %5d panels %dx%d cm\n" % (p2x1, 2*w, h) 
		txt += "  %5d panels %dx%d cm - " % (p1x2, w, 2*h) 
		txt += " %d left, %d right\n" % (p1x2_l, p1x2_r)
		txt += "  %5d panels %dx%d cm - " % (p1x1, w, h) 
		txt += " %d left, %d right\n" % (p1x1_l, p1x1_r)


		self.room_rep = {
			"txt":               txt,
			"area":              area,
			"active_area":       active_area,
			"panels_120x200":    p2x2,
			"panels_60x200":     p2x1,
			"panels_120x100":    p1x2,
			"panels_120x100_l":  p1x2_l,
			"panels_120x100_r":  p1x2_r,
			"panels_60x100":     p1x1,
			"panels_60x100_l":   p1x1_l,
			"panels_60x100_r":   p1x1_r,
		}

		return self.room_rep




	def draw_label(self, msp):

		write_text(msp, "Locale %d" % self.pindex, self.pos, zoom=2)


	def draw_lines(self, msp):

		for line in self.lines:
			line.draw_couplings(msp)
			line.draw_adductions(msp)

		#self.draw_connectors(msp)

	def draw_connectors(self, msp):
	
		arrng = self.arrangement
		if (not hasattr(arrng, 'alloc_clt_xside')):
			return
		cside = arrng.alloc_clt_xside
		cpos = self.collector.pos
		dst = MAX_DIST
		front = list()

		# attachment point
		for cir in arrng.circuits:
			front_cpl = cir.couplings[-1]
			fpos = front_cpl.orig1
			d = dist(fpos, cpos) 
			if d < dst:
				dst = d
				rpos = fpos
				cn = self.conn = front_cpl.pos
				self.thr = cn[1]

		# build front line
		tail = False
		for cir in arrng.circuits:
			fcir = cir.face()
			xcir = fcir[0][0]
			ycir0 = fcir[0][1]
			ycir1 = fcir[1][1]
			if tail:
				if ((xcir > lvl and cir.flip) or 
					(xcir < lvl and not cir.flip)):
					fcir[0] = xcir, front[-1][1]
				else:
					front[-1] = front[-1][0], ycir0

			front += fcir
			lvl = xcir
			tail = True

		self.front = front

		for line in filter(lambda x: not x.joined, self.lines):
			if not line.facing:
				line.level += 0.001

		self.lines.sort(key=lambda x: x.level)

		# determine line position
		for line in filter(lambda x: not x.joined, self.lines):
			if line.level >= self.thr:
				line.plvl = self.sup
				self.sup += 1
			else:
				self.inf += 1
				line.plvl = -self.inf
			
		# build paths to connectors
		fw = max(self.inf, self.sup)+1
		sgn = 1
		if (cside == LEFT):
			sgn = -1
			fw = -fw

		for k, line in enumerate(self.lines):
	
			if not hasattr(line, 'plvl'):
				continue

			# end of path
			k = line.plvl
			if k<0:
				k = -self.inf -k-1

			ew = cn[0] + fw*0.1, cn[1] + k*0.3 + 0.08
			ec = cn[0] + fw*0.1, cn[1] + k*0.3 - 0.08
	
			if not line.joined:
				
				ss = 1
				if line.plvl<0:
					ss = -1

				# start of path
				pos = line.couplings[-1].pos
				if line.facing:
					sw = pos[0], pos[1] + 0.08
					sc = pos[0], pos[1] - 0.08
				else:
					sw = pos[0]+sgn*0.2, pos[1] + 0.6 + 0.08
					sc = pos[0]+sgn*0.2, pos[1] + 0.6 - 0.08

				# find pipe level
				lvl = line.plvl
				if line.plvl <0:
					lvl = self.inf + lvl  
				lvl = lvl*0.05 + 0.1


				# build path
				polywarm = self.path(sw, ew, -sgn*lvl-sgn*ss*0.02)
				polycold = self.path(sc, ec, -sgn*lvl+sgn*ss*0.02)
				polywarm.append(self.abs((pos[0], pos[1] + 0.08)))
				polycold.append(self.abs((pos[0], pos[1] - 0.08)))
		
				pw = msp.add_lwpolyline(polywarm)
				pc = msp.add_lwpolyline(polycold)
				pw.dxf.color = color_warm
				pc.dxf.color = color_cold
				pw.dxf.layer = layer_link
				pc.dxf.layer = layer_link
			else:
				for l in line.lines:
					pos = l.couplings[-1].pos
					sw = pos[0], pos[1] + 0.08
					sc = pos[0], pos[1] - 0.08

					polywarm = self.path(sw, ew, lvl)
					polycold = self.path(sc, ec, lvl+0.02)
					pw = msp.add_lwpolyline(polywarm)
					pc = msp.add_lwpolyline(polycold)
					pw.dxf.color = color_warm
					pc.dxf.color = color_cold
					pw.dxf.layer = layer_link
					pc.dxf.layer = layer_link

	def path(self, s, e, offs):		
		global add_offs

		fc = deepcopy(self.front)
		fc = offset(fc, offs)	
		ps = project_hor(s, fc)
		pe = project_hor(e, fc)

		if ps[1] < pe[1]:
			fc = [p for p in fc if ps[1]<=p[1] and p[1]<=pe[1]]	
			fc = [s, ps] + fc + [pe, e]
			fc.reverse()
		else:
			fc = [p for p in fc if pe[1]<=p[1] and p[1]<=ps[1]]	
			fc = [e, pe] + fc + [ps, s]

		pfront = []
		for p in fc:
			pfront.append(self.abs(p))


		pfront = miter(pfront, add_offs)	
		return pfront

	def draw_passive(self, msp):
		global scale

		hatch = msp.add_hatch(color=41)
		if (frame_enabled):
			hatch.set_pattern_fill("ANSI31", scale=20)
		else:
			hatch.set_pattern_fill("ANSI31", scale=2/scale)

		hatch.paths.add_polyline_path(self.points, 
			is_closed=True,
			flags=ezdxf.const.BOUNDARY_PATH_EXTERNAL)

		hatch.dxf.layer = layer_lux

		for panel in self.panels:
			hatch.paths.add_polyline_path(panel.poly(), 
				is_closed=True,
				flags=ezdxf.const.BOUNDARY_PATH_OUTERMOST)

			# draw lux
			if panel.size[1] == 2:
				lux = msp.add_hatch(color=stripe_color)
				lux.dxf.layer = layer_lux
				if (frame_enabled):
					lux.set_pattern_fill("ANSI31", scale=20)
				else:
					lux.set_pattern_fill("ANSI31", scale=2/scale)
				luxp = panel.lux_poly()
				lux.paths.add_polyline_path(luxp, is_closed=True)
				pl = msp.add_lwpolyline(luxp)
				pl.dxf.layer = layer_lux	

		for obstacle in self.obstacles:
			hatch.paths.add_polyline_path(obstacle.points, 
				is_closed=True,
				flags=ezdxf.const.BOUNDARY_PATH_OUTERMOST)

	def draw_probes(self, msp):

		if len(self.panels) == 0:
			return

		if (self.color==bathroom_color
				and self.area_m2<=9):
			probe_type = "sonda T"

		if (self.color==valid_room_color or
		     (self.color==bathroom_color
				and self.area_m2>=9)):
			probe_type = "sonda T_U"

		x, y = self.panels[0].center()
		x = x - 25/scale
		y = y - 25/scale

		probe = msp.add_blockref(probe_type, (x,y), 
			dxfattribs={'xscale': 0.07/scale, 
   						'yscale': 0.07/scale})
		probe.dxf.layer = layer_probes
		

	def draw_structure(self, msp):

		global scale

		if len(self.panels) == 0:
			return

		p = self.panels[0].polyline()

		# horizontal versor
		norm = dist(p[2], p[1])
		vx = (p[2][0] - p[1][0])/norm
		vy = (p[2][1] - p[1][1])/norm

		# vertical versor
		norm = dist(p[1], p[0])
		ux = (p[1][0] - p[0][0])/norm
		uy = (p[1][1] - p[0][1])/norm

		minx = MAX_DIST
		maxx = -MAX_DIST
		miny = MAX_DIST
		maxy = -MAX_DIST

		for point in self.points:
			x, y = point[0], point[1]
			px, py = x*vx + y*vy, x*ux + y*uy
			minx = min(px, minx)
			maxx = max(px, maxx)
			miny = min(py, miny)
			maxy = max(py, maxy)
	
		miny -= 10/scale
		maxy += 10/scale 	
		s0 = sx0, sy0 = minx*vx + miny*ux, minx*vy + miny*uy
		q0 = qx0, qy0 = maxx*vx + miny*ux, maxx*vy + miny*uy
		s1 = sx1, sy1 = minx*vx + maxy*ux, minx*vy + maxy*uy
		q1 = qx1, qy1 = maxx*vx + maxy*ux, maxx*vy + maxy*uy
		hl = dist(s0, q0) 

		# projecting panels on axis
		xcoords = list()
		for panel in self.panels:
			p = panel.polyline()
			x0, y0 = p[1][0], p[1][1]
			x1, y1 = p[2][0], p[2][1]

			for n in range(2*panel.size[0]+1):
				ax  = vx * (x0 - sx0) + n*50/scale
				ax += vy * (y0 - sy0)
				xcoords.append(ax)
		
		xcoords.sort()

		# Complete structure left and right
		xc = list()
		m = xcoords[0] - 50/scale
		while(m>0):
			xc.append(m)
			m -= 50/scale
		
		m = xcoords[-1] + 50/scale
		while(m<hl):
			xc.append(m)
			m += 50/scale

		# Add structure if >60cm
		m = xcoords[0]
		xc.append(m)
		for x in xcoords:
			if (x-m < 1/scale):
				 continue

			xc.append(x)

			if (x-m>60/scale):
				n = int((x-m)/(60/scale))
				if n>0:
					for k in range(n):
						xc.append(m + (k+1)*(x-m)/n)
			m = x
			
		for ax in xc:
			a0 = (minx+ax)*vx + miny*ux, (minx+ax)*vy + miny*uy
			a1 = (minx+ax)*vx + maxy*ux, (minx+ax)*vy + maxy*uy
			css = trim(self.points, (a0, a1))
			for c in css:
				pl = msp.add_lwpolyline([c[0],c[1]])
				pl.dxf.layer = layer_struct
				pl.dxf.color = zone_color 
			
		#points = list()
		#for panel in self.panels:
		#	p = panel.polyline()
		#	norm = dist(p[2], p[1])
			# vx = (p[2][0] - p[1][0])/norm
			# vy = (p[2][1] - p[1][1])/norm
			# print("panel.size", panel.size[1])
			# for k in range(2*panel.size[1]+1):
			# 	print(k)

			# 	ax = p[0][0] + k*(50/scale)*vx
			# 	ay = p[0][1] + k*(50/scale)*vy
			# 	a = (ax, ay)
			# 	bx = p[1][0] + k*(50/scale)*vx
			# 	by = p[1][1] + k*(50/scale)*vy
			# 	b = (bx, by)

			# 	q0 = extend_to_poly((a, b), self.poly)
			# 	q1 = extend_to_poly((b, a), self.poly)

			# 	if not q0 or not q1:
			# 		continue

			# 	d = MAX_DIST
			# 	for pp in points:
			# 		d1 = min(dist(pp, q0), dist(pp, q1))
			# 		if d1 < d:
			# 			d = d1

			# 	if d<0.1/scale:
			# 		continue

			# 	print(q1, q0)

			# 	points.append(q1)
			# 	points.append(q0)
			# 	pline = [q1,q0]
			# 	pl = msp.add_lwpolyline(pline)
			# 	pl.dxf.layer = layer_struct
			# 	pl.dxf.color = zone_color
			# 	pl.dxf.linetype = 'CONTINUOUS'


	def draw(self, msp):
	
		if (debug):
			self.arrangement.draw_grid(msp)


		for panel in self.panels:
			panel.draw(msp)
			if (debug):
				panel.draw_profile(msp)

		self.draw_lines(msp)
		self.draw_label(msp)
		self.draw_passive(msp)
		self.draw_structure(msp)


class Model(threading.Thread):
	def __init__(self, output):
		super(Model, self).__init__()
		self.rooms = list()
		self.vectors = list()
		self.collectors = list()
		self.valid_rooms = list()
		self.processed = list()
		self.obstacles = list()
		self.zone = list()
		self.zone_bb = list()
		self.user_zones = list()
		self.output = output
		self.best_list = list()
		self.text_nav = ""
		self.cnd = list()
		self.laid = "without"

	def new_layer(self, layer_name, color):
		attr = {'linetype': 'CONTINUOUS', 'color': color}
		self.doc.layers.new(name=layer_name, dxfattribs=attr)
		

	def create_layers(self):
		self.new_layer(layer_panel, 0)
		if (debug):
			self.new_layer(layer_panelp, 0)
			self.new_layer(layer_box, box_color)
		self.new_layer(layer_link, 0)
		self.new_layer(layer_text, text_color)
		self.new_layer(layer_error, 0)
		self.new_layer(layer_lux, 0)
		self.new_layer(layer_probes, 0)
		self.new_layer(layer_joints, 0)
		self.new_layer(layer_struct, 0)

		self.doc.layers.get(layer_lux).off()
		self.doc.layers.get(layer_struct).off()

	def find_gates(self):
		
		for room1 in self.processed:
			for room2 in self.processed:
				if (room1 != room2):
					room1.add_gates(room2)

	def collector_side(self):

		for room in self.processed:
			if (not room.collector):
				self.output.print("CRITICAL: Room %d disconnected\n" % room.index)
				continue

			cltr = room.collector
			(vx, vy) = (cltr.pos[0]-room.pos[0], 
							cltr.pos[1]-room.pos[1])

			if (vx>=0):
				room.clt_xside = RIGHT
			else:
				room.clt_xside = LEFT

			if (vy>=0):
				room.clt_yside = TOP
			else:
				room.clt_yside = BOTTOM

	def route(self):

		# reverse path from rooms
		for room in self.processed:
			#print(room.links)
			pass


	def create_trees(self):

		# create trees
		self.find_gates()
		for room in self.processed:
			room.links = list()
			room.zone = False

		zone = 1
		for collector in self.collectors:
			collector.is_leader = False
			collector.zone_rooms = list()

			for room in self.processed:
				room.visited = False
				room.uplink = None
				room.walk = MAX_DIST

			root = collector.contained_in
			if (not root):
				#print("collector at ", collector.pos, " outside rooms")
				continue

			root.walk = 0
			root.uplink = root
			
			root.set_as_root(self.processed.copy(), collector)

			leader = None
			for room in self.processed:

				if room.user_zone != collector.user_zone:
					room.walk = MAX_DIST

				# check if room is assigned to a collector
				if ((not room.zone) and room.walk<MAX_DIST):
					# now assign  collector
					
					if (not collector.is_leader):
						collector.is_leader = True
						leader = collector
						collector.zone_num = zone
						collector.number = 1
						collector.next_item = 2
						collector.name = 'C' + str(zone) + '.1'
						zone += 1

					room.zone = collector
					leader = room.zone
					collector.zone_rooms.append(room)

				else:
					if (room.zone and room.walk<MAX_DIST):
						leader = room.zone
 
				link_item = (collector, room.walk, room.uplink)
				room.links.append(link_item)

			if (not collector.is_leader):
				if (leader):
					collector.zone_num = leader.zone_num
					collector.number = leader.next_item
					collector.name = 'C' + str(leader.zone_num)
					collector.name += '.' + str(leader.next_item)
					leader.next_item += 1

				else:
					collector.name ="unassigned"
					collector.zone_num = 0
					collector.number = 0
	
		# Assign rooms with fixed collectors to zones
		for room in self.rooms:
			if room.fixed_collector:
				collector = room.fixed_collector
				leader = collector.contained_in.zone 
				leader.zone_rooms.append(room)
				room.zone = leader

		self.best_dist = MAX_DIST
		for collector in self.collectors:
			collector.freespace = feeds_per_collector 
			collector.freeflow = flow_per_collector
			collector.items = list()
			for room in self.processed:
				if (room.fixed_collector and 
					room.fixed_collector == collector):
					collector.items.append(room)

		self.processed.sort(key=lambda x: x.links[0][1], reverse=True)

		# trim distance vectors
		for room in self.processed:
			room.links.sort(key=lambda x: x[1])
			if (room.links[0][1]> max_clt_distance 
					and not room.fixed_collector):
				self.output.print(
					"ABORT: No collectors from Room %d\n" % room.pindex)
				self.output.print("Check %s layer" % layer_error + 
					" to visualize errors\n")

				room.poly.dxf.layer = layer_error
				self.output_error()
				return False


		bound = 0
		for room in reversed(self.processed):
			if (room.fixed_collector or 
				room.color==disabled_room_color):
				room.bound = 0
				continue
			room.bound = bound
			bound += room.links[0][1]
			#print(room.pindex, room.bound, room.links[0][1])

		for room in self.processed:

			for i, link in enumerate(room.links):
				if (link[1]>max_clt_distance 
					or i>=max_clt_break):
					break

			if (i+1 < len(room.links)):
				del room.links[i:]

		return True

	def rescale_model(self):

		global tolerance 
		global font_size
		global default_scale
		global default_panel_width
		global default_panel_height
		global default_search_tol
		global default_hatch_width
		global default_hatch_height
		global default_collector_size
		global default_min_dist
		global default_min_dist2
		global default_wall_depth
		global default_max_clt_distance
		global default_add_offs
		global default_add_dist

		global panel_width 
		global panel_height 
		global search_tol 
		global hatch_width
		global hatch_height
		global collector_size
		global min_dist
		global min_dist2
		global wall_depth
		global max_clt_distance
		global add_offs
		global add_dist

		tolerance    = default_tolerance/scale
		font_size  = default_font_size/scale

		panel_width = default_panel_width/scale
		panel_height = default_panel_height/scale
		search_tol = default_search_tol/scale
		hatch_width = default_hatch_width/scale
		hatch_height = default_hatch_height/scale
		collector_size = default_collector_size/scale
		min_dist = default_min_dist/scale
		min_dist2 = default_min_dist2/scale
		wall_depth = default_wall_depth/scale
		max_clt_distance = default_max_clt_distance/scale
		add_offs = default_add_offs/scale
		add_dist = default_add_dist/scale

	def autoscale(self):
		global scale
		
		for e in self.msp.query('*[layer=="%s"]' % self.inputlayer):
			if (e.dxftype() == 'LINE'):
				continue
			if (e.dxftype() != 'LWPOLYLINE'):
				wstr = "WARNING: layer contains non-polyline: %s\n" % e.dxftype()
				self.Textinfo.insert(END, wstr)

		searchstr = 'LWPOLYLINE[layer=="'+self.inputlayer+'"]'
		query = self.msp.query(searchstr)
		if (len(query) == 0):
			wstr = "WARNING: layer %s does not contain polylines\n" % self.inputlayer
			self.output.print(wstr)

		n = 0
		tot = 0
		# Create list of rooms
		for poly in query:
			tot += Room(poly, self.output).area
			n += 1

		scale = pow(10, ceil(log10(sqrt(n/tot))))
		self.output.print("Autoscale: 1 unit = %g cm\n" % scale)

	def check_polyline_color(self, poly):

		if (poly.dxf.color == 256):
			# BYLAYER poly color
			poly.dxf.color = self.layer_color

		if (not (poly.dxf.color == collector_color or
				 poly.dxf.color == obstacle_color or
				 poly.dxf.color == bathroom_color or
				 poly.dxf.color == valid_room_color or
				 poly.dxf.color == zone_color or
				 poly.dxf.color == disabled_room_color)):
			return False
		return True


	def run(self):
 
		global tot_iterations, max_iterations
		global scale


		if self.refit:
			self.output.print("******************************************\n");
			self.output.print("Detected existing plan, disable allocation\n");
			self.output.print("******************************************\n");

		scale = self.scale
		if (self.scale == "auto"):
			scale = default_scale
			self.rescale_model()
			self.autoscale()
		else:
			scale = float(self.scale)

		self.rescale_model()

		Room.index = 1
		if not self.refit:
			self.create_layers()
		
		for e in self.msp.query('*[layer=="%s"]' % self.inputlayer):
			if (e.dxftype() == 'LINE'):
				self.vectors.append(e)
				continue

			if (e.dxftype() != 'LWPOLYLINE'):
				wstr = "WARNING: layer contains non-polyline: %s\n" % e.dxftype()
				self.textinfo.insert(END, wstr)

		searchstr = 'LWPOLYLINE[layer=="'+self.inputlayer+'"]'
		self.query = self.msp.query(searchstr)
		if (len(self.query) == 0):
			wstr = "WARNING: layer %s does not contain polylines\n" % self.inputlayer
			self.output.print(wstr)


		pindex = 1
		# Create list of rooms, obstacles and collectors
		for poly in self.query:

			# check if poly color is allowed
			if (not self.check_polyline_color(poly)):
				wstr = "ABORT: Polyline color %d not allowed" % poly.dxf.color
				self.output.print(wstr)
				return

			room = Room(poly, self.output)
			self.rooms.append(room)

			if (len(room.errorstr)>0):
				# Invalid polyline
				room.error = True
				self.textinfo.insert(END,room.errorstr)
			else:

				# Valid polyline, classify room
				room.error = False
				area = scale * scale * room.area
				if (area > max_room_area and
				    (room.color == valid_room_color or
				     room.color == bathroom_color)):
					wstr = "ABORT: Zone %d larger than %d m2\n" % (room.index, 
						max_room_area)
					wstr += "Consider splitting area \n\n"
					self.output.print(wstr)
					room.errorstr = wstr
					room.error = True
					continue

				if (room.color == collector_color):
					self.collectors.append(room)
					room.is_collector = True
					continue

				if (room.color == valid_room_color or
				   room.color == bathroom_color):
					self.valid_rooms.append(room)

				if (room.color == obstacle_color):
					self.obstacles.append(room)

				if (room.color == disabled_room_color):
					room.pindex = pindex
					pindex += 1
					self.processed.append(room)

				if (room.color == zone_color):
					self.user_zones.append(room)

	
		# check if the room is too small to be processed
		for room in self.valid_rooms:
			area = scale * scale * room.area
			if  (area < min_room_area):
				wstr = "WARNING: area less than %d m2: " % min_room_area
				wstr += "Consider changing scale!\n"
				self.output.print(wstr)
				room.errorstr = wstr
				room.error = True
			else:
				room.pindex = pindex
				pindex += 1
				self.processed.append(room)


		# check if every collector is in a room
		for collector in self.collectors:

			collector.contained_in = False
			for room in self.processed:

				if (room.contains(collector)):
					collector.contained_in = room
					room.obstacles.append(collector)
					break

			if (not collector.contained_in):
				wstr = "ABORT: Collector outside room"
				self.output.print(wstr)
				return
			
		# assings collectors to user zone
		for collector in self.collectors:
			collector.user_zone = None
			for zone in self.user_zones:
				if zone.contains(collector):
					if collector.user_zone != None:
						wstr = "ABORT: Collector inside two user zones"
						self.output.print(wstr)
						return
					else:
						collector.user_zone = zone

		# assign rooms to user zones
		for room in self.processed:
			room.user_zone = None
			for zone in self.user_zones:
				if zone.contains(room):
					if room.user_zone != None:
						wstr = "ABORT: Room inside two user zones"
						self.output.print(wstr)
						return
					else:
						room.user_zone = zone
					

		# assign obstacles to rooms
		for obs in self.obstacles:
			for room in self.processed:
				if (obs.collides_with(room)):
					room.obstacles.append(obs)
					obs.contained_in = room

		# check if two rooms collide
		self.valid_rooms.sort(key=lambda room: room.ax)	
		room = self.valid_rooms
		for i in range(len(room)):
			j=i+1
			while (j<len(room) and room[j].ax < room[i].bx):
				if (room[i].collides_with(room[j])):
					wstr = "ABORT: Collision between Room %d" % room[i].pindex
					wstr += " and Room %d \n" % room[j].pindex
					wstr += ("Check %s in output drawing" % layer_error +
					 " to visualize errors")
					room[i].poly.dxf.layer = layer_error
					room[j].poly.dxf.layer = layer_error
					self.output.print(wstr)
					self.output_error()
					return

					#if (room[i].area > room[j].area):
					#	room[i].obstacles.append(room[j])
					#	room[j].contained_in = room[i]
					#else:
					#	room[j].obstacles.append(room[i])
					#	room[i].contained_in = room[j]
				j += 1

		# check if two collectors collide
		for i in range(len(self.collectors)-1):
			for j in range(i+1, len(self.collectors)):
				if (self.collectors[i].collides_with(self.collectors[j])):
					wstr = "ABORT: Collision between collectors\n"
					wstr += ("Check %s layer " % layer_error +
					 "to visualize errors")
					self.collectors[i].poly.dxf.layer = layer_error
					self.collectors[j].poly.dxf.layer = layer_error
					self.output.print(wstr)
					self.output_error()
					return
	
		# check if vector is in room or 
		# across collector and room
		for v in self.vectors:

			p1 = (v.dxf.start[0], v.dxf.start[1])
			p2 = (v.dxf.end[0], v.dxf.end[1])

			p1_clt = p2_clt = None
			for clt in self.collectors:
				if clt.is_point_inside(p1):
					p1_clt = clt
					break;

				if clt.is_point_inside(p2):
					p2_clt = clt
					break

			for room in self.processed:

				if p1_clt and room.is_point_inside(p2):
					room.fixed_collector = p1_clt
					self.msp.delete_entity(v)
					break

				if p2_clt and room.is_point_inside(p1):
					room.fixed_collector = p2_clt
					self.msp.delete_entity(v)
					break

				if (room.contains_vector(v) and
					not (p1_clt or p2_clt)):
					# Allocate vector
					room.vector = v
					norm = dist(p1, p2)
					uv = room.uvector = (p2[0]-p1[0])/norm, (p2[1]-p1[1])/norm
					room.rot_orig = p1
					room.rot_angle = -atan2(uv[1], -uv[0])*180/pi
					break
			else:
				# Check if vector is vector fixes to collector

				wstr = "ABORT: Vector outside room\n"
				wstr += ("Check %s layer" % layer_error + 
					" to visualize errors")
				v.dxf.layer = layer_error
				self.output.print(wstr)
				self.output_error()
				return
	
		# orient room without vector
		for room in self.processed:
			if (not room.vector):
				room.orient_room()	
	
		self.output.print("Detected %d rooms\n" % len(self.processed))
		self.output.print("Detected %d collectors\n" % len(self.collectors))

		# Check if room too large  for a collector
		for room in self.processed:
			area = scale * scale * room.area
			flow = area*flow_per_m2
			if flow > flow_per_collector:
				wstr = "ABORT: Room %d larger than collector capacity\n" % room.pindex
				wstr += ("Check %s layer" % layer_error + 
					" to visualize errors\n")
				room.poly.dxf.layer = layer_error
				self.output.print(wstr)
				self.output_error()
				return

		# Check if enough collectors
		tot_area = feeds_eff = feeds_max = 0
		flow_eff = flow_max = 0
		for room in self.processed:
			area = scale * scale * room.area
			room.feeds_eff = ceil(area/area_per_feed_m2*target_eff)
			room.feeds_max = ceil(area/area_per_feed_m2)
			feeds_eff += room.feeds_eff
			feeds_max += room.feeds_max
			# connect room based on max allocation
			room.feeds = room.feeds_max
			room.flow_eff = area*target_eff*flow_per_m2
			room.flow_max = area*flow_per_m2
			flow_eff += room.flow_eff
			flow_max += room.flow_max
			room.flow = room.flow_eff
			#self.output.print("Room%3d   lines:%2d  flow:%6.2lf l/h\n" % 
			#	(room.pindex, room.feeds, room.flow_eff))
			tot_area += area

		available_feeds = feeds_per_collector * len(self.collectors)
		available_flow  = flow_per_collector * len(self.collectors)
		self.output.print("Available pipes %g\n" % available_feeds)
		self.output.print("Estimated pipes for %d%% cover: %d\n" % 
				(100*target_eff, feeds_eff))
		self.output.print("Estimated pipes for 100%% cover: %d\n" % feeds_max)
		self.output.print("Available flow %g l/h\n" % available_flow)
		self.output.print("Estimated flow for %d%% cover: %d l/h\n" % 
				(100*target_eff, flow_eff))
		self.output.print("Estimated flow for 100%% cover: %d l/h\n" % flow_max)
		if (feeds_eff > available_feeds or flow_eff > available_flow):
			self.output.print("ABORT: Too few collectors\n")
			self.output.print("Please insert at least %d collectors\n" %
				ceil(flow_eff/flow_per_collector))
			return

		if (feeds_max > available_feeds or flow_max > available_flow):
			rc = ceil(flow_eff/flow_per_collector)
			if rc>len(self.collectors):
				self.output.print("WARNING: Possible insufficient collectors\n")
				self.output.print("WARNING: suggested %d collectors\n" % rc)

		################################################################
		if not self.create_trees():
			return

		# Disabled room with collector forms its own zone
		for room in self.processed:
			if (room.color==disabled_room_color and room.zone and
				not room.collector):
					room.collector = room.zone	

		# for collector in self.collectors:
		#self.draw_trees(self.collectors[5])

		#self.draw_gates()	
		#self.doc.saveas(self.outname)
		#return

		################################################################
		#  Mapping rooms to collectors

		self.processed.append(None)    ;# Add sentinel
		room_iter = iter(self.processed)
		tot_iterations = 0
		self.found_one = False
		self.output.print("Linking Rooms:")
		self.connect_rooms(room_iter, 0)
		self.output.print("\n")
		self.processed.pop()           ;# Remove sentinel
		if (not self.found_one):
			self.output.print("CRITICAL: Could not connect rooms\n")
			return

		#self.draw_uplinks()
		#self.draw_trees(self.collectors[3])

		if not self.refit:
			self.populating_model()

		##############################################################
		# summary
		# self.output.clear()
		report = self.print_report()

		################################################################
		# Splitting/joining circuits 

		fline = self.ptype['flow_line']
		fpanel = self.ptype['flow_panel']
		for room in self.processed:
			room.make_lines(fline, fpanel)

		self.joints = 0
		for room in self.processed:
			cirs = room.arrangement.circuits
			tot_split = 0
			for cir in cirs:
				cir.flow = cir.size * self.ptype['flow_panel']
				cir.split = False
				if (cir.flow>self.ptype['flow_line']):
					tot_split += 1
					cir.split = True

			room.joints = len(cirs) + tot_split - room.actual_feeds
			self.joints += room.joints

		################################################################

		summary = ""
		if (not self.mtype == "none"):
			summary += self.output_html()
		summary += '<div class="section">'
		summary += '<h4>Relazione calcolo</h4>'
		summary += '<pre>'
		summary += report
		summary += '</pre></div>'

		if (not web_version):
			self.output.print(summary)
		out = self.outname[:-4]+".txt"
		f = open(out, "w")
		print(summary, file = f)

		#if (web_version):
		#	slink = os.path.dirname(out) + "/output.txt"
		#	if (os.path.exists(slink)):
		#		os.remove(slink)
		#	os.symlink(out, slink)

		print("SUMMARY DONE")

		##############################################################

		if not self.refit:
			self.draw()
		else:
			self.doc.saveas(self.outname)



		##############################################################
		# save data in XLS

		self.save_in_xls()
		self.save_navision()
		self.thumbnail()

		if not self.refit and self.laid != "without":
			self.save_in_word()

		print("ALL DONE")

	def populating_model(self):

		# Determine which side is collector in each room
		self.collector_side()

		# allocating panels in room	
		self.output.print("Processing Room:")
		self.processed.sort(key = lambda x: x.pindex)
		count = 5
		for room in self.processed:
			if (room.color == disabled_room_color):
				continue

			self.output.print("%d " % room.pindex)
			room.alloc_panels()
		
			count += 1
			if (count == 19):
				self.output.print("\n")
				count = 0

		self.output.print("\n")

		################################################################

		# find fittings 
		for room in self.processed:
			if (hasattr(room.arrangement, 'alloc_clt_xside')):
				room.arrangement.fittings()
				room.arrangement.make_couplings()
	
		# find attachment points of dorsals
		self.dorsals = list()
		for room in self.processed:
			room_dors = room.arrangement.dorsals
			mode = room.arrangement.alloc_mode
			self.dorsals += room_dors
			uplink_dist = MAX_DIST
			for dorsal in room_dors:
				# dorsal.cltrs = list()
				if (mode==0):
					dorsal.pos = dorsal.attach
				else:
					dorsal.pos = (dorsal.attach[1], dorsal.attach[0])

				d = dist(dorsal.pos, room.uplink.pos)
				if (d <= uplink_dist):
					uplink_dist = d
					room.attachment = dorsal.pos

	def output_error(self):
		self.doc.layers.remove(layer_panel)
		self.doc.layers.remove(layer_link)
		if (debug):
			self.doc.layers.remove(layer_box)
			self.doc.layers.remove(layer_panelp)

		for room in self.processed:
			room.draw_label(self.msp)

		if (os.path.isfile(self.outname) and ask_for_write==True):
			if askyesno("Warning", "File 'leo' already exists: Overwrite?"):
				self.doc.saveas(self.outname)
		else:
			self.doc.saveas(self.outname)

		#if (web_version):
		#	slink = os.path.dirname(self.outname) + "/output.dxf"
		#	if (os.path.exists(slink)):
		#		os.remove(slink)
		#	os.symlink(self.outname, slink)

		print("DRAW DONE")

	def output_html(self):

		if (self.mtype == "none"):
			return ""

		mtype = self.mtype[:-5]
		mtype_label = ac_label[mtype]
		mount = 'V' if self.mtype[-4:] == "vert" else 'O'
		mount_label = 'verticale' if self.mtype[-4:] == "vert" else 'orizzontale'

		for ac in air_conditioners:
			if (mtype == ac['type'] and mount == ac['mount']):
				self.cnd.append(ac)

		html =  '<div class="section" >'
		html += '<h4>Eurotherm consiglia</h4>'

		for clt in self.collectors:
			if not clt.is_leader:
				continue
				
			html += '<p id="suggest"><b><u>Zona %d</u></b></p>' % clt.zone_num

			zone_area = 0
			for room in clt.zone_rooms:
				zone_area += room.area_m2

			self.volume = float(self.height) * zone_area

			html += '<p id="suggest">%s ad installazione %s ' % (mtype_label, mount_label)
			html += 'per una portata di %.2f m3/h:</p>' % self.volume

			self.find_air_conditioners()
			self.zone.append(self.best_ac)

			for k, ac in enumerate(self.cnd):
				if (self.best_ac[k] > 0):
					html += '<p id="mtype">%d x ' % self.best_ac[k] + ac['model'] + '</p>'
			html += '<p id="suggest">copertura %.2f m3/h,' % self.best_flow
			html += 'eccesso %.2f m3/h</p>' % (self.best_flow-self.volume)

		# final count 
		l = len(self.cnd)
		self.best_ac = [0]*l
		for zone in self.zone:
			for k in range(l):
				self.best_ac[k] += zone[k]

		html += '</div>'
		return html


	def find_air_conditioners(self):
	
		l = len(self.cnd)
		self.cnd.sort(key= lambda x: x["flow_m3h"]);
		max_ac = ceil(self.volume/self.cnd[0]["flow_m3h"])
		num_tot = l*max_ac
		
		num_ac = [0]*l
		self.best_ac = [0]*l
		self.best_flow = MAX_COST
		for i in range((max_ac+1) ** l):
			flowtot = 0
			count = i
			for k in range(l):
				val = count % (max_ac+1)
				count = count//(max_ac+1)
				num_ac[k] = val
				flowtot += val * self.cnd[k]["flow_m3h"]
			if (flowtot >= self.volume and flowtot < self.best_flow):
				self.best_flow = flowtot
				for k, val in enumerate(num_ac):
					self.best_ac[k] = num_ac[k]


	def draw(self):
		global collector_size, search_tol

		# Box zones
		for clt in self.collectors:
			if (not clt.is_leader):
				continue

			# Box zones
			if not clt.user_zone:	
				ax = min([c.ax for c in clt.zone_rooms]) - 2*min_dist
				ay = min([c.ay for c in clt.zone_rooms]) - 2*min_dist
				bx = max([c.bx for c in clt.zone_rooms]) + 2*min_dist
				by = max([c.by for c in clt.zone_rooms]) + 2*min_dist
			
				pline = [(ax,ay),(ax,by),(bx,by),(bx,ay),(ax,ay)]
				pl = self.msp.add_lwpolyline(pline)
				pl.dxf.layer = layer_text
				pl.dxf.color = zone_color
				pl.dxf.linetype = 'CONTINUOUS'
			else:
				ax = -MAX_DIST
				by = -MAX_DIST

				for p in clt.user_zone.poly:
					if p[1] > by or (p[1] == by and p[0] < ax):
						ax = p[0]; by = p[1]

			write_text(self.msp, "Zone %d" % clt.zone_num, (ax, min_dist+by), 
				align=ezdxf.lldxf.const.MTEXT_BOTTOM_LEFT)
			self.zone_bb.append([ax,ay,bx,by])


		# Collectors
		for collector, items in self.best_list:
			feeds = 0
			for room in items:
				feeds += room.actual_feeds
			collector.req_feeds = feeds
			collector.label = " (%d+%d)" % (feeds, feeds)
			txt = collector.name + collector.label
			xc, yc = collector.pos[0], collector.pos[1]
			cs = collector_size

			write_text(self.msp, txt, (xc-cs/2, yc+cs/2+search_tol), 
				align=ezdxf.lldxf.const.MTEXT_BOTTOM_LEFT,
				zoom=0.6,
				col=collector_color)
			xs, ys = 0.1/scale, 0.1/scale
			orig = xc - cs/2, yc - cs/2

			block = self.msp.add_blockref(block_collector, orig, 
				dxfattribs={'xscale': xs, 'yscale': ys})

			block.dxf.layer = layer_panel

		for room in self.processed:
			room.draw(self.msp)
			if self.control == "reg":
				room.draw_probes(self.msp)


		## drawing connections
		self.draw_links2()

		#for collector in self.collectors:
		#	self.draw_trees(collector)
		# self.draw_gates()

		##############################################################

		if (frame_enabled):
			# Rescale everything before printing	
			block = self.doc.blocks.new("DRAWING")
			ss = list()
			for e in self.msp:
				if not e.dxf.layer == "Cornice":
					block.add_entity(e)
					ss.append(e)


			for e in ss:
				self.msp.unlink_entity(e)

			# bounding box zones	
			ax = self.zone_bb[0][0]
			ay = self.zone_bb[0][1]
			bx = self.zone_bb[0][2]
			by = self.zone_bb[0][3]

			for zone in self.zone_bb:
				ax = min(ax, zone[0])
				ay = min(ay, zone[1])
				bx = max(bx, zone[2])
				by = max(by, zone[3])

			ax = ax*scale*10
			ay = ay*scale*10
			bx = bx*scale*10
			by = by*scale*10

			# centering
			if bx-ax>a1_bx-a1_ax or by-ay>a1_by-a1_ay:
				cx = (a0_bx + a0_ax)/2
				cy = (a0_by + a0_ay)/2
			else:
				cx = (a1_bx + a1_ax)/2
				cy = (a1_by + a1_ay)/2

			cx -= (bx + ax)/2
			cy -= (by + ay)/2

			orig = (cx, cy) 
			block1 = self.msp.add_blockref("DRAWING", orig, 
				dxfattribs={'xscale': scale*10, 'yscale': scale*10})

			block1.explode()

		##############################################################

		self.doc.layers.get(layer_lux).off()

		if (debug):
			self.doc.layers.get(layer_box).off()
			self.doc.layers.get(layer_panelp).off()

		if (os.path.isfile(self.outname) and ask_for_write==True):
			if askyesno("Warning", "File 'leo' already exists: Overwrite?"):
				self.doc.saveas(self.outname)
		else:
			self.doc.saveas(self.outname)

		#if (web_version):
		#	slink = os.path.dirname(self.outname) + "/output.dxf"
		#	if (os.path.exists(slink)):
		#		os.remove(slink)
		#	os.symlink(self.outname, slink)

		print("DRAW DONE")


	def draw_links(self):	
		# draw connections
		for collector, items in self.best_list:
			for room in items:
				if (len(room.panels) == 0):
					continue

				pos = room.attachment
				while (room != collector.contained_in):
					(p1, p2) = room.wall(room.uplink)
					med = ((p1[0]+p2[0])/2, (p1[1]+p2[1])/2)
					pline = (pos, med)
					pl = self.msp.add_lwpolyline(pline)
					pl.dxf.layer = layer_panel
					pl.dxf.color = 4
					pos = med
					room = room.uplink
				
				pline = (pos, collector.pos)
				pl = self.msp.add_lwpolyline(pline)
				pl.dxf.layer = layer_panel
				pl.dxf.color = 4

	def draw_links2(self):	
		# draw connections

		for collector, items in self.best_list:

			for room in items:
				if (len(room.panels) == 0):
					continue

				time_to_live = 16

				pos = room.pos
				while (room != collector.contained_in):
					time_to_live = time_to_live - 1
					if (time_to_live == 0):
						print("Uplink error in room", room.index)
						break

					# med = room.pos
					# pline = (pos, med)
					# pl = self.msp.add_lwpolyline(pline)
					# pl.dxf.layer = layer_link
					# pl.dxf.color = 6
					# pos = med
					if (room.uplink == room):
						break
					room = room.uplink
				
				#pline = (pos, collector.pos)
				#pl = self.msp.add_lwpolyline(pline)
				#pl.dxf.layer = layer_link
				#pl.dxf.color = 6
				#pl.dxf.lineweight = 2

				dx = collector.pos[0] - pos[0]
				dy = collector.pos[1] - pos[1]
				d = dist(pos, collector.pos)
				ux, uy = -dy/d, dx/d

				sx = pos[0] + 2/scale*ux
				sy = pos[1] + 2/scale*uy
				ex = collector.pos[0] + 2/scale*ux
				ey = collector.pos[1] + 2/scale*uy
				pline = ((ex, ey), (sx, sy))				
				pl = self.msp.add_lwpolyline(pline)
				pl.dxf.layer = layer_link
				pl.dxf.color = color_warm
				pl.dxf.lineweight = 50

				sx = pos[0] - 2/scale*ux
				sy = pos[1] - 2/scale*uy
				ex = collector.pos[0] - 2/scale*ux
				ey = collector.pos[1] - 2/scale*uy
				pline = ((ex, ey), (sx, sy))				
				pl = self.msp.add_lwpolyline(pline)
				pl.dxf.layer = layer_link
				pl.dxf.color = color_cold
				pl.dxf.lineweight = 50


	def draw_trees(self, collector):

		for room in self.processed:

			uplink = None
			for item in room.links:
				if (collector == item[0]):
					cltr, walk, uplink = item
					break

			if (not uplink):
				continue

			(x0, y0) = room.pos
			pl = self.msp.add_circle(room.pos,2*search_tol)
			pl.dxf.layer = layer_link
			pl.dxf.color = 1

			(x1, y1) = uplink.pos
			pline = ((x0,y0), (x1,y1))
			pl = self.msp.add_lwpolyline(pline)
			pl.dxf.layer = layer_link
			pl.dxf.color = 1
			
			if (not room.uplink):
				continue

			pl = self.msp.add_circle(room.uplink.pos,2*search_tol)
			pl.dxf.layer = layer_link
			pl.dxf.color = 1

	def draw_uplinks(self):

		for room in self.processed:

			if (not room.uplink):
				continue

			(x0, y0) = room.pos
			pl = self.msp.add_circle(room.pos,2*search_tol)
			pl.dxf.layer = layer_link
			pl.dxf.color = 1

			(x1, y1) = room.uplink.pos
			pline = ((x0,y0), (x1,y1))
			pl = self.msp.add_lwpolyline(pline)
			pl.dxf.layer = layer_link
			pl.dxf.color = 1
			pl = self.msp.add_circle(room.uplink.pos,2*search_tol)
			pl.dxf.layer = layer_link
			pl.dxf.color = 1


	def draw_gates(self):

		for room in self.processed:
			for newroom, wall in room.gates:
				pl = self.msp.add_lwpolyline(wall)
				pl.dxf.layer = layer_link
				pl.dxf.color = 6
				pline = (room.pos, newroom.pos)
				pl = self.msp.add_lwpolyline(pline)
				pl.dxf.layer = layer_link
				pl.dxf.color = 5

				pl = self.msp.add_circle(room.pos,4*search_tol)
				pl.dxf.layer = layer_link
				pl.dxf.color = 5

	# Connect rooms to collectors using branch-and-bound
	def connect_rooms(self, room_iter, partial):

		global tot_iterations, max_iterations

		# Check if time to give up
		if (tot_iterations % 100e3 == 0):
			self.output.print("#")
		tot_iterations += 1
		if (tot_iterations > max_iterations):
			return

		room = next(room_iter)
		while room and len(room.links)==0:
			room = next(room_iter)
			 
		# Terminal case
		if (room == None):

			if (partial < self.best_dist):
				# Found solution
				self.found_one = True

				self.best_dist = partial
				self.best_list = list()
				ir = iter(self.processed)
				while (x:=next(ir)) != None:
					if x.color == disabled_room_color:
						continue
					if x.fixed_collector:
						x.downlinks = list()
						x.collector = x.fixed_collector
						x.uplink = x.collector.contained_in
						continue
					x.downlinks = list()
					x.uplink = x._uplink
					x.collector = x._collector

				for collector in self.collectors:
					for room in collector.items:
						if (room.color == disabled_room_color
							or room.fixed_collector):
							continue
						room.uplink.downlinks.append(room)
					item = (collector, copy(collector.items))
					self.best_list.append(item)
				return

		# skip disabled room
		if room.color == disabled_room_color:
			self.connect_rooms(copy(room_iter), partial)
			return

		# Recursive cases
		for link in room.links:
			collector, room_dist, uplink = link
		
			# If room if fixed to a collector, skip
			# other collectors
			if room.fixed_collector:
				collector = room.fixed_collector
				room_dist = 0
				
			new_partial = partial + room_dist
			
			if ((new_partial+room.bound<self.best_dist and 
				collector.freespace>=room.feeds and
				collector.freeflow>=room.flow)):
				collector.items.append(room)
				collector.freespace -= room.feeds
				collector.freeflow  -= room.flow
				room._uplink = uplink
				room._collector = collector
				self.connect_rooms(copy(room_iter), new_partial)
				collector.items.remove(room)
				collector.freespace += room.feeds
				collector.freeflow += room.flow
					
	def print_report(self):
		
		txt = "\n ------- Room Report ----------\n\n"

		self.area = 0 
		self.active_area = 0 

		p2x2 = 0; p2x2_h = 0
		p2x1 = 0; p2x1_h = 0
		p1x2 = 0; p1x2_h = 0
		p1x2_l = 0; p1x2_h_l = 0
		p1x2_r = 0; p1x2_h_r = 0
		p1x1 = 0; p1x1_h = 0
		p1x1_l = 0; p1x1_h_l = 0
		p1x1_r = 0; p1x1_h_r = 0

		w = default_panel_width
		h = default_panel_height
		failed_rooms = 0
		self.feeds = 0
		self.normal_area = 0
		self.normal_active_area = 0
		self.normal_passive_area = 0
		self.bathroom_area = 0
		self.bathroom_active_area = 0
		self.bathroom_passive_area = 0

		for room in self.processed:

			if (len(room.errorstr)>0):
				failed_rooms += 1
				continue

			txt += "Room %d  --------- \n" % room.pindex
			if not self.refit:
				rep = room.report()
			else:
				rep = room.recount_panels(self.msp)

			self.area += rep['area']
			self.active_area += rep['active_area']
			txt += rep['txt'] + "\n"

			if (room.color == bathroom_color):
				self.bathroom_area += rep['area']
				self.bathroom_active_area += rep['active_area'] 
				self.bathroom_passive_area += rep['area'] - rep['active_area']
			else:
				self.normal_area += rep['area']
				self.normal_active_area += rep['active_area']
				self.normal_passive_area += rep['area'] - rep['active_area']

			if (room.color == valid_room_color):
				p2x2 += rep['panels_120x200']
				p2x1 += rep['panels_60x200']

				p1x2 += rep['panels_120x100']
				p1x1 += rep['panels_60x100']

				p1x1_l += rep['panels_60x100_l']
				p1x1_r += rep['panels_60x100_r']
				p1x2_l += rep['panels_120x100_l']
				p1x2_r += rep['panels_120x100_r']
			else:
				p2x2_h += rep['panels_120x200']
				p2x1_h += rep['panels_60x200']

				p1x2_h += rep['panels_120x100']
				p1x1_h += rep['panels_60x100']

				p1x1_h_l += rep['panels_60x100_l']
				p1x1_h_r += rep['panels_60x100_r']
				p1x2_h_l += rep['panels_120x100_l']
				p1x2_h_r += rep['panels_120x100_r']
				

			room.actual_feeds = ceil(rep['active_area']/area_per_feed_m2)
			self.feeds += room.actual_feeds
			
		self.passive_area = self.area - self.active_area
		self.perimeter = 0
		for room in self.processed:
			self.perimeter += room.perimeter
			
		# Summary of all areas
		smtxt =  "\n\nTotal processed rooms %d\n" % len(self.processed)
		smtxt += "Total collectors %d\n" % len(self.collectors)
		smtxt += "Total area %.2lf m2\n" % self.area
		smtxt += "Total active area %.2lf m2 " % self.active_area
		smtxt += " (%.2lf %%)\n" % (100*self.active_area/self.area)
		smtxt += "Total passive area %.2lf m2\n" % self.passive_area
		smtxt += "Normal area %.2lf m2\n" % self.normal_area
		smtxt += "Normal active area %.2lf m2\n" % self.normal_active_area
		smtxt += "Normal passive area %.2lf m2\n" % self.normal_passive_area
		smtxt += "Hydro area %.2g m2\n" % self.bathroom_area
		smtxt += "Hydro active area %.2lf m2\n" % self.bathroom_active_area
		smtxt += "Hydro passive area %.2lf m2\n" % self.bathroom_passive_area
		smtxt += "Total perimeter %.2lf m\n" % (self.perimeter*scale/100)
		smtxt += "Total pipes %d\n" % self.feeds
		smtxt += "Normal panels count:\n"
		smtxt += "  %5d panels %dx%d cm\n" % (p2x2, 2*w, 2*h) 
		smtxt += "  %5d panels %dx%d cm\n" % (p2x1, 2*w, h) 
		smtxt += "  %5d panels %dx%d cm - " % (p1x2, w, 2*h)
		smtxt += "  %d left, %d right\n" % (p1x2_l, p1x2_r)
		smtxt += "  %5d panels %dx%d cm - " % (p1x1, w, h) 
		smtxt += "  %d left, %d right\n" % (p1x1_l, p1x1_r)
		smtxt += "Hydro panels count:\n"
		smtxt += "  %5d panels %dx%d cm\n" % (p2x2_h, 2*w, 2*h) 
		smtxt += "  %5d panels %dx%d cm\n" % (p2x1_h, 2*w, h) 
		smtxt += "  %5d panels %dx%d cm - " % (p1x2_h, w, 2*h)
		smtxt += "  %d left, %d right\n" % (p1x2_h_l, p1x2_h_r)
		smtxt += "  %5d panels %dx%d cm - " % (p1x1_h, w, h) 
		smtxt += "  %d left, %d right\n" % (p1x1_h_l, p1x1_h_r)

		self.laid_half_panels   = 2*(p2x2   + p2x1)   + p1x2   + p1x1
		self.laid_half_panels_h = 2*(p2x2_h + p2x1_h) + p1x2_h + p1x1_h

		# Requirements normal panels
		smtxt += "\n> Requirements:\n"
		p2x2_cut = min(p1x2_r,p1x2_l) + abs(p1x2_r-p1x2_l)
		self.panels_120x200 = p2x2_tot = p2x2 + p2x2_cut 
		p2x2_spr = abs(p1x2_r-p1x2_l)
		smtxt += "  %d panels %dx%d, \n" % (p2x2_tot, 2*w, 2*h)
		smtxt += "    of which %d to cut and %d halves spares\n" % (p2x2_cut, p2x2_spr)
		p2x1_cut = min(p1x1_r,p1x1_l) + abs(p1x1_r-p1x1_l)
		self.panels_60x200 = p2x1_tot = p2x1 + p2x1_cut 
		p1x1_spr = abs(p1x1_r-p1x1_l)
		smtxt += "  %d panels %dx%d, \n" % (p2x1_tot, 2*w, 2*h)
		smtxt += "    of which %d to cut and %d halves spares\n" % (p2x1_cut, p1x1_spr) 

		# Requirements waterproof panels
		smtxt += "\n> Requirements Hydro:\n"
		p2x2_h_cut = min(p1x2_h_r,p1x2_h_l) + abs(p1x2_h_r-p1x2_h_l)
		self.panels_h_120x200 = p2x2_h_tot = p2x2_h + p2x2_h_cut 
		p2x2_h_spr = abs(p1x2_h_r-p1x2_h_l)
		smtxt += "  %d panels %dx%d, \n" % (p2x2_h_tot, 2*w, 2*h)
		smtxt += "    of which %d to cut and %d halves spares\n" % (p2x2_h_cut, p2x2_h_spr)
		p2x1_h_cut = min(p1x1_h_r,p1x1_h_l) + abs(p1x1_h_r-p1x1_h_l)
		self.panels_h_60x200 = p2x1_h_tot = p2x1_h + p2x1_h_cut 
		p1x1_h_spr = abs(p1x1_h_r-p1x1_h_l)
		smtxt += "  %d panels %dx%d, \n" % (p2x1_h_tot, 2*w, 2*h)
		smtxt += "    of which %d to cut and %d halves spares\n" % (p2x1_h_cut, p1x1_h_spr) 

		return smtxt + txt



	def save_in_xls(self):
		wb = openpyxl.load_workbook(xlsx_template)
		ws1 = wb[sheet_template_1]
		ws2 = wb[sheet_template_2]
		ws3 = wb[sheet_template_3]
	
		sheet_bd = deepcopy(sheet_breakdown)
		sheet_bd[0] += (ws1,)
		sheet_bd[1] += (ws2,)
		sheet_bd[2] += (ws3,)

		no_collectors = len(self.collectors)

		thin_left = Border(left=Side(style='thin'))
		thin_right = Border(right=Side(style='thin'))
		thin_left_top = Border(left=Side(style='thin'),
						top=Side(style='thin'))
		thin_right_top = Border(right=Side(style='thin'),
						top=Side(style='thin'))

		# copy total area
		ws1['B14'] = self.area
		ws2['B14'] = self.area
		ws3['B14'] = self.area
		ws1['B18'] = self.area
		ws2['B18'] = self.area
		ws3['B18'] = self.area

		# copy total coverage
		ws1['A14'] = self.active_area
		ws2['A14'] = self.active_area
		ws3['A14'] = self.active_area
		ws1['A18'] = self.active_area
		ws2['A18'] = self.active_area
		ws3['A18'] = self.active_area

		# copy total panels
		ws1['D14'] = ws1['D18'] = self.panels_120x200 + self.panels_h_120x200
		ws1['D14'] = ws1['D18'] = self.panels_120x200 + self.panels_h_120x200
		ws2['D14'] = ws2['D18'] = self.panels_120x200 + self.panels_h_120x200
		ws2['D14'] = ws2['D18'] = self.panels_120x200 + self.panels_h_120x200
		ws3['D14'] = ws3['D18'] = self.panels_120x200 + self.panels_h_120x200
		ws3['D14'] = ws3['D18'] = self.panels_120x200 + self.panels_h_120x200

		ws1['E14'] = ws1['E18'] = self.panels_60x200 + self.panels_h_60x200
		ws1['E14'] = ws1['E18'] = self.panels_60x200 + self.panels_h_60x200
		ws2['E14'] = ws2['E18'] = self.panels_60x200 + self.panels_h_60x200
		ws2['E14'] = ws2['E18'] = self.panels_60x200 + self.panels_h_60x200
		ws3['E14'] = ws3['E18'] = self.panels_60x200 + self.panels_h_60x200
		ws3['E14'] = ws3['E18'] = self.panels_60x200 + self.panels_h_60x200

		# copy number of feeds
		ws1['F14'] = ws1['F18'] = self.feeds
		ws1['F14'] = ws1['F18'] = self.feeds
		ws2['F14'] = ws2['F18'] = self.feeds
		ws2['F14'] = ws2['F18'] = self.feeds
		ws3['F14'] = ws3['F18'] = self.feeds
		ws3['F14'] = ws3['F18'] = self.feeds

		# copy number of collectors 
		ws1['G14'] = ws1['G18'] = no_collectors 
		ws1['G14'] = ws1['G18'] = no_collectors 
		ws2['G14'] = ws2['G18'] = no_collectors 
		ws2['G14'] = ws2['G18'] = no_collectors 
		ws3['G14'] = ws3['G18'] = no_collectors 
		ws3['G14'] = ws3['G18'] = no_collectors 


		for sheet, warm_coef, cool_coef, wsh in sheet_bd:
			#ws = wb.create_sheet(sheet)
			ws = wsh

			ws['A22'] = sheet
			ws.row_dimensions[3].height = 32

			# header
			for i in range(65,85):
				ws.column_dimensions[chr(i)].width = 10
				ws[chr(i)+'23'].alignment = \
					Alignment(wrapText=True, 
						vertical ='center',
						horizontal ='center')

			ws['A23'] = "Zona"
			ws['B23'] = "Collettore"
			ws['C23'] = "Stanza"

			ws['D23'] = "Attiva\n[m2]"
			ws['E23'] = "Area\n[m2]"
			ws['F23'] = "% cop."
			ws['G23'] = "linee"

			ws['H23'] = "Pannelli\n200x120"
			ws['I23'] = "Pannelli\n200x60"
			ws['J23'] = "Pannelli\n100x120"
			ws['K23'] = "Pannelli\n100x60"
				
			#ws.column_dimensions['M'].width = 2

			ws['L22'] = 'Riscaldamento'
			ws['L23'] = "Q, resa\n[W]"
			ws['M23'] = "Q, tot\n[W]"
			ws['N23'] = "Portata\n[kg/h]"

			#ws.column_dimensions['Q'].width = 2

			ws['O22'] = 'Raffrescamento'
			ws['O23'] = "Q, resa\n[W]"
			ws['P23'] = "Q, tot\n[W]"
			ws['Q23'] = "Portata\n[kg/h]"

			set_border(ws, '23', "ABCDEFGHIJKLMNOPQ")
			ws['A23'].border = thin_left_top
			ws['B23'].border = thin_left_top
			ws['G23'].border = thin_left_top
			ws['L23'].border = thin_left_top
			ws['O23'].border = thin_left_top
			ws['Q23'].border = thin_right_top

			self.processed.sort(key=lambda x: 
				(x.collector.zone_num, x.collector.number, x.pindex))

			zone = 0
			index = 24
			number = -1
			for room in self.processed:

				ws['A'+str(index)].border = thin_left
				ws['B'+str(index)].border = thin_left
				ws['G'+str(index)].border = thin_left
				ws['L'+str(index)].border = thin_left
				ws['O'+str(index)].border = thin_left
				ws['Q'+str(index)].border = thin_right

				# body
				for i in range(65,82):
					pos = chr(i)+str(index)
					ws[pos].alignment = \
						Alignment(wrapText=True, 
							vertical ='center',
							horizontal ='center')
					color = "D0D0D0"
					if (index%2==0 and i>65):
						color = "F0F0F0"
					ws[pos].fill = PatternFill(start_color=color, 
							fill_type = "solid")

				while (room.collector.zone_num>zone):
					zone += 1
					pos = 'A' + str(index)
					ws[pos] = "Zona %d" % zone
					set_border(ws,str(index), 'B')
				
				if (room.collector.number != number):
					number = room.collector.number
					set_border(ws, str(index), "ABCDEFGHIJKLMNOPQ")
					ws['A'+str(index)].border = thin_left_top
					ws['B'+str(index)].border = thin_left_top
					ws['G'+str(index)].border = thin_left_top
					ws['L'+str(index)].border = thin_left_top
					ws['O'+str(index)].border = thin_left_top
					ws['Q'+str(index)].border = thin_right_top


				pos = 'B' + str(index)
				ws[pos] = room.collector.name
				ws[pos].alignment = Alignment(horizontal='center')
				
				pos = 'C' + str(index)
				ws[pos] = room.pindex
				ws[pos].alignment = Alignment(horizontal='center')

				pos = 'E' + str(index)
				ws[pos] = room.area_m2
				ws[pos].number_format = "0.0"

				if (room.active_m2==0):
					index += 1
					continue

				pos = 'D' + str(index)
				ws[pos] = room.active_m2
				ws[pos].number_format = "0.0"

				pos = 'F' + str(index)
				ws[pos] = room.ratio
				ws[pos].number_format = "0.0"

				pos = 'G' + str(index)
				ws[pos] = room.actual_feeds

				if (room.room_rep['panels_120x200']>0):
					pos = 'H' + str(index)
					ws[pos] = room.room_rep['panels_120x200']

				if (room.room_rep['panels_60x200']>0):
					pos = 'I' + str(index)
					ws[pos] = room.room_rep['panels_60x200']

				if (room.room_rep['panels_120x100']>0):
					pos = 'J' + str(index)
					ws[pos] = room.room_rep['panels_120x100']

				if (room.room_rep['panels_60x100']>0):
					pos = 'K' + str(index)
					ws[pos] = room.room_rep['panels_60x100']


				# heating 
				pos = 'L' + str(index)
				ws[pos] = radiated = room.active_m2 * warm_coef
				ws[pos].number_format = "0"

				pos = 'M' + str(index)
				ws[pos] = output = radiated * 1.1
				ws[pos].number_format = "0"

				pos = 'N' + str(index)
				ws[pos] = 3.6*output/(4.186*wsh['K14'].value)
				ws[pos].number_format = "0"

				# cooling
				pos = 'O' + str(index)
				ws[pos] = absorbed = room.active_m2 * cool_coef
				ws[pos].number_format = "0"

				pos = 'P' + str(index)
				ws[pos] = output = absorbed * 1.1
				ws[pos].number_format = "0"

				pos = 'Q' + str(index)
				ws[pos] = 3.6*output/(4.186*wsh['K18'].value)
				ws[pos].number_format = "0"


				#ws[pos_area].number_format = "0.00"
				index += 1

			set_border(ws, str(index), "ABCDEFGHIJKLMNOPQ")


		if (web_version):
			out = self.outname[:-4] + ".xlsx"
		else:
			out = self.filename[:-4] + ".xlsx"	
		
		wb.save(out)

		#if (web_version):
		#	slink = os.path.dirname(out) + "/output.xlsx"
		#	if (os.path.exists(slink)):
		#		os.remove(slink)
		#	os.symlink(out, slink)

	def save_navision(self):
		global scale

		# Radiant panels
		self.text_nav += nav_item(self.panels_120x200*2.4, 
				self.ptype['code_full'], self.ptype['desc_full'])

		self.text_nav += nav_item(self.panels_60x200*1.2,
				self.ptype['code_half'], self.ptype['desc_half'])

		self.text_nav += nav_item(self.panels_h_120x200*2.4, 
				self.ptype['code_full_h'], self.ptype['desc_full_h'])

		self.text_nav += nav_item(self.panels_h_60x200*1.2, 
				self.ptype['code_half_h'], self.ptype['desc_half_h'])

		# Passive panels
		code = '6111020101'
		desc = 'LEONARDO PASSIVO 1200x2000x50mm'
		qnt = 1.05*self.normal_passive_area
		self.text_nav += nav_item(qnt, code, desc)
		code = '6114020201'
		desc = 'LEONARDO PASSIVO IDRO 1200x2000x50mm'
		qnt = 1.05*self.bathroom_passive_area
		self.text_nav += nav_item(qnt, code, desc)

		# Couplings
		for fit in fittings:
			(fittings[fit])['count'] = 0 

		for e in self.msp:	
			if e.dxftype() == "INSERT":
				name = e.block().name
				if name[0:3] == "Rac":
					tag = name.split("_")[-1] 
					if tag=="rosso" or tag=="blu":
						fit = "_".join(name.split("_")[:-1])
					else:
						fit = name
					fittings[fit]['count'] += 1

		fittings['Rac_20_20_20']['count'] = 2*self.joints
		
		# count the plugs == circuits
		plugs = (fittings['Rac_20_10']['count'] +
				fittings['Rac_20_10_10']['count'] +
				fittings['Rac_10_20_10']['count'] +
				fittings['Rac_10_20_10_10']['count'] +
				fittings['Rac_10_10_20_10_10']['count'])//2

		for name in fittings:
			fit = fittings[name]
			if fit['count']:
				desc = "RACCORDO LEONARDO " + fit['desc']
				self.text_nav += nav_item(fit['count'],
					fit['code'], desc)

		# bent joint.
		if not self.refit:
			code = '6910022006'
			desc = 'CURVA LEONARDO 20-20 (4pz)'
			qnt1 = ceil(0.04*self.area)
			self.text_nav += nav_item(qnt1, code, desc)


		# Interconnection pipes #############################

		# use collector labels to count collectors and circuits
		clt_qnts = [0]*(feeds_per_collector+1)
		tot_cirs = 0
		tot_clts = 0
		for e in self.msp.query('*[layer=="%s"]' % layer_text):	
			if e.dxftype() == "MTEXT" and e.text[0] == 'C':
				tot_clts += 1
				tag = e.text.split()[1][1:-1]
				feeds = int(tag.split("+")[0])

				if (feeds==0 and not self.refit):
					self.output.print(
						"WARNING: Collector %s unused\n" % e.text[:-6])

				if (feeds<=0 and self.refit):
					self.output.print(
						"ABORT: Label %s not recognized as a collector name\n" % e.text)
					self.output_error()
					return

				if (feeds>=feeds_per_collector):
					self.output.print(
						"ABORT: Too many lines in collector %s\n" % e.text)
					self.output_error()
					return
					
				clt_qnts[feeds] += 1
				tot_cirs += feeds

		code = '2112200220'
		desc = 'TUBO MULTISTRATO 20X2 RIV.ROSSO'
		qnt = self.area * 0.7
		self.text_nav += nav_item(qnt, code, desc)
		code = '2112200120'
		desc = 'TUBO MULTISTRATO 20X2 RIV.BLU'
		self.text_nav += nav_item(qnt, code, desc)

		code = '2720200120'
		desc = 'LINEA AGG. PERT-AL-PERT + ANELLI E TERMIN. (2m)'
		qnt = 0
		for e in self.msp.query('*[layer=="%s"]' % layer_panel):
			if e.dxftype() == "LWPOLYLINE":
				points = list(e.vertices())	
				mdist = 0
				for i in range(len(points)-1):
					mdist = max(mdist, dist(points[i], points[i+1]))
				mdist = int(np.round(mdist*scale/100))
				qnt += mdist

		qnt = ceil(qnt/2)
		self.text_nav += nav_item(qnt, code, desc)

		# linear joint
		# code = '6910022005'
		# desc = 'RACCORDO LEONARDO 20-20 (4pz)'
		# self.text_nav += nav_item(qnt2, code, desc)

		# rings
		code = '6910022011'
		desc = 'ANELLO PER RACC.LEONARDO IN PLASTICA D20 (8pz)'	
		qnt1 = fittings['Rac_20_20_dritto']['count']
		qnt2 = fittings['Rac_20_20_curva']['count']
		qnt = 2*tot_cirs + 3*2*self.joints + 2*qnt1 + 2*qnt2
		self.text_nav += nav_item(qnt, code, desc)


		# Control panel
		closures = 0
		if (self.ptype['handler']=='30'):
			code = '6113021002'
			desc = 'LEONARDO QUADRO DI CHIUSURA PLUS'
			qnt = ceil(0.25*(self.laid_half_panels+self.laid_half_panels_h))
			closures += qnt
			self.text_nav += nav_item(qnt, code, desc)
		else:
			code = '6110020103'
			desc = 'LEONARDO QUADRO CHIUSURA RACCORDI 420x260mm'
			qnt = ceil(0.25*self.laid_half_panels)
			closures += qnt
			self.text_nav += nav_item(qnt, code, desc)

			code = '6112020201'
			desc = 'LEONARDO QUADRO CHIUSURA RACC.AMB.UMIDI 420x260mm'
			qnt = ceil(0.25*self.laid_half_panels_h)
			closures += qnt
			self.text_nav += nav_item(qnt, code, desc)

		code = '6920042001'
		desc = 'COLLA PER QUADRI DI CHIUSURA'
		qnt = ceil(closures/4.35)
		self.text_nav += nav_item(qnt, code, desc)


		tot_adpt = 0
		for i in range(1,feeds_per_collector+1):
			tot_adpt += clt_qnts[i]*i
			if (clt_qnts[i] == 0):
				continue
			code = '41200101%02d' % i
			desc = 'COLLETTORE SL 1" %02d+%02d COMPLETO' % (i,i)
			self.text_nav += nav_item(clt_qnts[i],code, desc)	
		
		# Adaptors
		code = '4810202001'
		desc = 'ADATTATORE'
		qnt = 2*tot_adpt
		self.text_nav += nav_item(qnt,code, desc)

		# Hatch (botola)
		code = '6920012001'
		desc = 'BOTOLA'
		self.text_nav += nav_item(tot_clts, code, desc)
		code = '4910020112'
		desc = 'MANOMETRO'
		self.text_nav += nav_item(tot_clts, code, desc)
		code = '4710020306'
		desc = 'COPPIA VALVOLE SFERA SL SQ/DR.1"1/4 F - 1"F'
		self.text_nav += nav_item(tot_clts, code, desc)
		code = '4713010301'
		desc = 'ISOLAZIONE VALVOLA SFERA SL (coppia)'
		self.text_nav += nav_item(tot_clts, code, desc)

		# headers (testina)	
		if not self.control == "reg":
			code = '5150020201'
			desc = 'TESTINE ELETTROTERMICHE 4 FILI'
			self.text_nav += nav_item(tot_cirs, code, desc)
		#code = '5150020202'
		#desc = 'TESTINE 2 FILI'
		#self.text_nav += nav_item(tot_cirs, code, desc)

		fpanel = self.ptype['flow_panel']
		zones = list()
		smartp = 0
		smartp_b = 0
		for room in self.processed:
			flow = fpanel*room.active_m2/2.4

			if room.color == disabled_room_color:
				continue

			if not room.zone in zones:
				zones.append(room.zone)
				room.zone.zone_count = 1
				room.zone.flow = flow
				room.zone.smartp = 0
				room.zone.smartp_b = 0
			else:
				room.zone.zone_count += 1
				room.zone.flow += flow

			if (self.mtype == "none" or 
				room.color == bathroom_color 
				and room.area_m2<=9):
				room.zone.smartp_b += 1
				smartp_b += 1
			else:
				room.zone.smartp += 1
				smartp += 1

		# if regulated
		if self.control == "reg":

			smartp = 0
			smartp_b = 0
			for e in self.msp.query('*[layer=="%s"]' % layer_probes):	
				if e.block().name == "sonda T_U":
					smartp += 1
				else:
					smartp_b += 1
				

			smartbases = 0
			smartcomforts = 0
			for zone in zones:
				zone.smartbases = ceil((zone.smartp+zone.smartp_b)/8)
				smartbases += zone.smartbases
				smartcomforts += ceil(zone.smartbases/8)


			code = '5150020202'
			desc = 'TESTINE ELETTROTERMICHE 2 FILI'
			qnt = tot_cirs
			self.text_nav += nav_item(qnt, code, desc)
		
			code = '5140030101'
			desc = 'SMARTCOMFORT 365'
			qnt = smartcomforts
			self.text_nav += nav_item(qnt, code, desc)

			code = '5140020401'
			desc = 'SMARTPOINT TEMPERATURA'
			qnt = smartp_b
			self.text_nav += nav_item(qnt, code, desc)

			code = '5140020402'
			desc = 'SMARTPOINT TEMPERATURA / UMIDITA\''
			qnt = smartp
			self.text_nav += nav_item(qnt, code, desc)

			code = '5140020201'
			desc = 'SMARTBASE PER LA GESTIONE DI TESTINE, POMPA E MISCELATRICE'
			qnt = smartbases
			self.text_nav += nav_item(qnt, code, desc)


			code = '5140020403'
			desc = 'SMARTPOINT EXT SONDA ESTERNA'
			qnt = smartcomforts
			self.text_nav += nav_item(qnt, code, desc)

			code = '5140020404'
			desc = 'SONDA DI MANDATA'
			qnt = smartcomforts 
			self.text_nav += nav_item(qnt, code, desc)


		# If air conditioning
		if not self.mtype == "none":
			compamat_R = 0
			compamat_TOP = 0
			compamat_SUPER = 0
			for zone in zones:
				if zone.flow < 1850:
					compamat_R += 1
				else:
					if zone.flow < 4000:
						compamat_TOP += 1
					else:
						compamat_SUPER +=1

			mtype = self.mtype[:-5]
			smartairs = 0
			if mtype=="dehum_int":
				smartairs = len(zones)

			if compamat_R > 0:
				code = '5330010101'
				desc = 'COMPAMAT R'
				qnt = compamat_R
				self.text_nav += nav_item(qnt, code, desc)

			if compamat_TOP > 0:
				code = '5330010201'
				desc = 'COMPAMAT TOP'
				qnt = compamat_TOP
				self.text_nav += nav_item(qnt, code, desc)

			if compamat_SUPER > 0:
				code = '5330010301'
				desc = 'COMPAMAT SUPER'
				qnt = compamat_SUPER
				self.text_nav += nav_item(qnt, code, desc)

			code = '5140020202'
			desc = 'SMARTAIR PER LA GESTIONE DELL\'UNITA\' ARIA E SERRANDE'
			qnt = smartairs
			self.text_nav += nav_item(qnt, code, desc)

			if self.control == "reg":
				code = '5140020301'
				desc = 'SET DI CONNETTORI SMARTBASE / SMARTAIR'
				qnt = smartbases + smartairs
				self.text_nav += nav_item(qnt, code, desc)

			# COMPAMAT accessories
			for k, cnd in enumerate(self.cnd):
				if (self.best_ac[k]>0):
					code = cnd['code']
					desc = cnd['model']
					qnt = self.best_ac[k]
					self.text_nav += nav_item(qnt, code, desc)
					accs = cnd['accessories'] 
					for acc in accs:
						if type(acc) == tuple:
							code = accessories[acc[0]]['code']
							desc = accessories[acc[0]]['desc']
							num = acc[1]
						else: 
							code = accessories[acc]['code']
							desc = accessories[acc]['desc']
							num = 1
							
						self.text_nav += nav_item(qnt*num, code, desc)
						
		# Abdution lines
		# Red stripes

		# corrosion inhibitor
		code = '3310020201'
		desc = 'e100 INIBITORE'
		if self.ptype['handler'] == '55':
			water = self.active_area*0.043*16.67 + self.area*0.41

		if self.ptype['handler'] == '35':
			water = self.active_area*0.043*23.34 + self.area*0.41
		
		if self.ptype['handler'] == '30':
			water = self.active_area*0.043*27.50 + self.area*0.41

		self.text_nav += nav_item(ceil(water/100), code, desc)	


		if (web_version):
			out = self.outname[:-4] + ".dat"
		else:
			out = self.filename[:-4] + ".dat"	
	
		f = open(out, "w")
		print(self.text_nav, file = f, end="")
		f.close()


	def entry(self, row, val):

		self.document.writetab(2, (row,5), Word.getv(val))
		price = self.document.gettab(2, (row,6))
		val = Word.euro(val*price)
		self.document.writetab(2, (row,7), val)	


	def save_in_word(self):

		global scale

		self.document = doc = Word("quote.docx")
		document = doc.document

		today = Word.date()

		doc.write(44,'Data: ' + today)
		doc.write(45,'Cliente: ' + self.cname)
		doc.write(46,'Rif. Cantiere: ' + self.caddr)
		doc.write(47,'Commessa: no. Ref.: ' + self.ccomp)

		#######################################################
		# table 1
		tt = document.tables[0].cell(1,2).text
		eur = document.tables[0].cell(1,4).text[0]

		unit = tt[1:]
		mq = "%.1f" % self.area + unit

		# first line
		doc.writetab(0, (1,2), mq)
		v1 = self.area * 57
		num = Word.euro(v1)
		doc.writetab(0, (1,4), num)

		# second line
		doc.writetab(0, (2,2), mq)
		v2 = self.area * 15
		num = Word.euro(v2) 
		doc.writetab(0, (2,4), num)
		
		# total
		num = Word.euro(v1+v2) 
		doc.writetab(1, (0,1), num)

		#######################################################		
		# table 2 
		obs_small = obs_medium = 0
		for room in self.processed:
			for obs in room.obstacles:
				obs_area =  obs.area*10000*scale*scale
				if obs_area<75:
					obs_small += 1 

				if 75 <= obs_area<= 400:
					obs_medium += 1 


		self.entry(4, obs_medium)
		self.entry(6, obs_small)
		self.entry(10, self.area)

		#ext11

		#ext12
		self.entry(12, self.perimeter*scale/100)

		# insert thumbnail
		img = self.outname[:-4] + ".png"
		document.paragraphs[132].alignment = 1
		r = document.paragraphs[132].add_run()
		r.add_picture(img, width=docx.shared.Inches(6.0))

		if (web_version):
			out = self.outname[:-4] + ".doc"
		else:
			out = self.filename[:-4] + ".doc"	

		document.save(out)
		print("DOCUMENT SAVED")	


	def thumbnail(self):
		# pictures of design
		os.system("python3 dxf2img.py "+self.outname+" > /dev/null")
		print("THUMBNAIL DONE")



class App:

	def __init__(self):
		self.loaded = False
		self.queue = queue.Queue()
		self.model = None

		self.root = Tk()
		root = self.root
		#root.geometry('500x300')
		root.title("Eurotherm Leonardo Planner - Experimental")
		root.resizable(width=False, height=False)

		# Control section
		ctlname = Label(root, text="Control")
		ctlname.grid(row=0, column=0, padx=(25,0), pady=(10,0), sticky="w")

		ctl = Frame(root)
		self.ctl = ctl
		ctl.config(borderwidth=1, relief='ridge')
		ctl.grid(row=1, column=0, padx=(25,25), pady=(0,20))

		button = Button(ctl, text="Open", width=5, command=self.search)
		button.grid(row=1, column=1, sticky="e")

		self.text = StringVar()
		flabel = Label(ctl, textvariable=self.text, width=30, anchor="w")
		flabel.config(borderwidth=1, relief='solid')
		flabel.grid(row=1, column=0, padx=(10,30), pady=(20,10))

		self.text1 = StringVar()
		flabel = Label(ctl, textvariable=self.text1, width=30, anchor="w")
		flabel.config(borderwidth=1, relief='solid')
		flabel.grid(row=2, column=0, padx=(10,30), pady=(10,20))

		self.var = StringVar() # variable for select layer menu

		self.button1 = Button(ctl, text="Build", width=5, command=self.create_model, pady=5)
		self.button1.grid(row=4, column=1, pady=(30,10))

		self.type = StringVar()
		names = [panel['full_name'] for panel in panel_types]
		self.type.set(names[0])
		self.typemenu = OptionMenu(self.ctl, self.type,*names)
		self.typemenu.config(width=26)
		self.typemenu.grid(row=4, column=0, padx=(10,40), sticky="w")

		# Parameters section
		parname = Label(root, text="Settings")
		parname.grid(row=2, column=0, padx=(25,0), pady=(1,0), sticky="w")
		self.params = params = Frame(root)
		params.config(borderwidth=1, relief='ridge')
		params.grid(row=3, column=0, sticky="ew", padx=(25,25), pady=(0,2))

		Label(params, text="A drawing unit in cm").grid(row=0, column=0, sticky="w")
		self.entry1 = Entry(params, justify='right', width=10)
		self.entry1.grid(row=0, column=1, sticky="w")
		self.entry1.insert(END, "auto")

		#Label(params, text="zone cost (m2)").grid(row=1, column=0, sticky="w")
		#self.entry2 = Entry(params, justify='right', width=10)
		#self.entry2.grid(row=1, column=1)
		#self.entry2.insert(END, str(default_zone_cost))

		#Label(params, text="transversal cuts").grid(row=2, column=0, sticky="e")
		#self.tcuts = IntVar()
		#self.entry3 = Checkbutton(params, variable=self.tcuts)
		#self.entry3.grid(row=2, column=1)

		# Info Section
		ctlname = Label(root, text="Report")
		ctlname.grid(row=4, column=0, padx=(25,0), pady=(10,0), sticky="w")

		self.textinfo = Text(root, height=20, width=58)
		self.textinfo.config(borderwidth=1, relief='ridge')
		self.textinfo.grid(row=5, column=0, pady=(0,15)) 
		#sb = Scrollbar(root, command=self.textinfo.yview)
		#sb.grid(row=3, column=1, sticky="nsw")


		self.reset()
		self.root.mainloop()

	def print(self, text):
		self.textinfo.insert(END, text)

	def clear(self):
		self.textinfo.delete('1.0', END)

	def search(self,event=None):
		self.filename = filedialog.askopenfilename(filetypes=[("DXF files", "*.dxf")])
		self.loadfile()

	def reset(self):
		self.text.set("Select DXF")
		self.text1.set("Modified DXF")
		self.button1["state"] = "disabled"
		self.textinfo.delete('1.0', END)

		if (hasattr(self,'opt')): self.opt.destroy()	
		self.var.set("Select layer")
		self.opt = OptionMenu(self.ctl, self.var,['0'])
		self.opt.config(width=26)
		self.opt.grid(row=3, column=0, padx=(10,40), sticky="w")

	def loadfile(self):
		try:
			self.doc_src = ezdxf.readfile(self.filename)
		except IOError:
			self.reset()
			return
		except ezdxf.DXFStructureError:
			self.textinfo.insert(END, 'Invalid or corrupted DXF file.')
			self.reset()
			return

		self.loaded = True

		self.text.set(self.filename)
		self.outname = self.filename[:-4]+"_leo.dxf"
		self.text1.set(self.outname)

		layers = [layer.dxf.name for layer in self.doc_src.layers]
		sel = layers[0]
		for layer in layers:
			if (layer == default_input_layer):
				sel = default_input_layer
				break

		self.opt.destroy()
		self.var.set(sel)
		self.opt = OptionMenu(self.ctl,self.var,*layers)
		self.opt.config(width=26)
		self.opt.grid(row=3, column=0, padx=(10,40), sticky="w")
		self.button1["state"] = "normal" 


	def create_model(self):

		if self.model and self.model.is_alive():
			return

		self.textinfo.delete('1.0', END)

		if (not self.loaded):
			self.textinfo(END, "File not loaded")
			return

		self.scale = self.entry1.get()
		self.inputlayer = self.var.get()
		self.mtype = "none"
		self.height = 2.7
		self.control = "reg"

		_create_model(self)


def _create_model(iface):

	global block_blue_120x100 
	global block_blue_60x100 
	global block_green_120x100
	global block_green_60x100
	global block_collector
	global area_per_feed_m2

	# create model and initialise it
	iface.model = Model(iface)

	# reload file
	iface.doc = ezdxf.readfile(iface.filename)	
	iface.model.refit = False
	for layer in iface.doc.layers:
		if layer.dxf.name == layer_panel:
			iface.model.refit = True

	if not iface.model.refit:	
		if (frame_enabled):
			iface.model.doc = ezdxf.readfile("frame3.dxf")
		else:
			iface.model.doc = ezdxf.new(dxf_version)
	else:
		iface.model.doc = iface.doc

	iface.model.doc.header["$LWDISPLAY"] = 1
	# iface.model.doc.header["$LWDISPSCALE"] = 0.55
	# self.model.doc = self.doc     # <<<<<<<<< MODIFIED LINE <<<<<<<

	iface.model.msp = iface.model.doc.modelspace()
	iface.model.scale = iface.scale
	iface.model.inputlayer = iface.inputlayer
	iface.model.textinfo = iface.textinfo
	iface.model.outname = iface.outname
	iface.model.filename = iface.filename
	iface.model.control = iface.control

	iface.model.mtype = iface.mtype
	iface.model.height = iface.height

	if web_version and iface.laid=="with":
		iface.model.laid = "with"
		iface.model.cname = iface.cname
		iface.model.caddr = iface.caddr
		iface.model.ccomp = iface.ccomp

	if iface.model.refit:
		if (not web_version):
			ctype = iface.type.get()
		else:
			ctype = iface.type
		for ptype in panel_types:
			if (ctype == ptype['full_name']):
				iface.model.ptype = ptype
		iface.model.start()
		return


	# copy input layer from source
	importer = Importer(iface.doc, iface.model.doc)
	ents = iface.doc.modelspace().query('*[layer=="%s"]' 
			% iface.model.inputlayer)

	for layer in iface.doc.layers:
		if (layer.dxf.name == iface.model.inputlayer):
			iface.model.layer_color = layer.dxf.color

	if (len(ents) == 0):
		iface.textinfo.print('Layer "%s" not available or empty'
			% iface.inputlayer)
		return

	importer.import_entities(ents)
	importer.finalize()

	## copy blocks from panels
	source_dxf = ezdxf.readfile("Symbol.dxf")
	importer = Importer(source_dxf, iface.model.doc)

	if (not web_version):
		ctype = iface.type.get()
	else:
		ctype = iface.type
	
	for ptype in panel_types:
		if (ctype == ptype['full_name']):
			iface.model.ptype = ptype
			handler = "LEO_" + ptype['handler'] + "_"
			block_blue_120x100 = handler + "120"
			block_blue_60x100 = handler + "60"
			block_green_120x100 = handler + "120_IDRO"
			block_green_60x100 = handler + "60_IDRO"

			if handler == "LEO_30_":
				block_green_120x100 = block_blue_120x100
				block_green_60x100 = block_blue_60x100

			area_per_feed_m2 = ptype['panels'] * 2.4
			flow_per_m2 = ptype['flow_panel'] / 2.4
			iface.print('Area/line = %g m2\n' % area_per_feed_m2)
			iface.print('Flow_per_m2 = %g l/m2\n' % flow_per_m2)


	importer.import_block(block_blue_120x100)
	importer.import_block(block_blue_60x100)
	importer.import_block(block_green_120x100)
	importer.import_block(block_green_60x100)
	importer.import_block(block_collector)
	importer.import_block("LEO_LUX_120")
	importer.import_block("LEO_LUX_120_IDRO")

	# import fittings
	importer.import_block("Rac_20_10_20_blu")
	importer.import_block("Rac_20_10_20_rosso")
	importer.import_block("Rac_20_10_blu")
	importer.import_block("Rac_20_10_rosso")
	importer.import_block("Rac_20_10_10_20_blu")
	importer.import_block("Rac_20_10_10_20_rosso")
	importer.import_block("Rac_20_10_10_blu")
	importer.import_block("Rac_20_10_10_rosso")
	importer.import_block("Rac_20_10_20_10_blu")
	importer.import_block("Rac_20_10_20_10_rosso")
	importer.import_block("Rac_10_20_10_blu")
	importer.import_block("Rac_10_20_10_rosso")
	importer.import_block("Rac_10_20_10_10_blu")
	importer.import_block("Rac_10_20_10_10_rosso")
	importer.import_block("Rac_20_10_10_20_10_10_blu")
	importer.import_block("Rac_20_10_10_20_10_10_rosso")
	importer.import_block("Rac_10_10_20_10_10_blu")
	importer.import_block("Rac_10_10_20_10_10_rosso")
	importer.import_block("Rac_20_20_dritto")
	importer.import_block("Rac_20_20_curva")
	importer.import_block("Rac_20_20_20_blu")
	importer.import_block("Rac_20_20_20_rosso")
	importer.import_block("Rac_10_10_dritto")
	importer.import_block("sonda T")
	importer.import_block("sonda T_U")
	importer.finalize()

	iface.model.start()

	
class Iface:
	def __init__(self, infile, units, ptype, control, 
	  laid, cname, caddr, ccomp,
	  mtype, height):
		self.filename = web_filename
		self.scale = units
		self.control = control
		self.mtype = mtype
		self.height = height
		self.inputlayer = default_input_layer
		self.textinfo = self
		self.outname = infile

		self.laid = laid
		self.cname = cname
		self.caddr = caddr
		self.ccomp = ccomp

		for	ctype in panel_types:
			if (ctype['handler'] == ptype):
				self.type = ctype['full_name']

		_create_model(self)

	def print(self, text):
		print(text, end='')


	def insert(self, pos, text):
		print(text, end='')

if (web_version):

	import atexit

	local_dir = os.path.dirname(os.path.realpath(__file__))
	os.chdir(local_dir)
	sys.path.append(local_dir + "/www/cgi-bin")

	from conf import *

	def remove_lock():
		os.remove(lock_name)
	
	if not os.path.exists(lock_name):

		# Acquire lock 
		open(lock_name, "w")	
		atexit.register(remove_lock)

		# Get command line parameters
		filename = sys.argv[1] 
		units = sys.argv[2]	
		ptype = sys.argv[3]
		control = sys.argv[4]

		laid = sys.argv[5]
		cname = sys.argv[6]
		caddr = sys.argv[7]
		ccomp = sys.argv[8]

		mtype = sys.argv[9]
		height = sys.argv[10]

		os.rename(filename, web_filename)
	
		# print("filename", filename)
		# print("units", units)
		# print("ptype", ptype)
		# print("control", control)
		# print("laid", laid)
		# print("cname", cname)
		# print("caddr", caddr)
		# print("ccomp", ccomp)
		# print("mtype", mtype)
		# print("height", height)
	
		Iface(filename, units, ptype, control, 
			laid, cname, caddr, ccomp,
			mtype, height)
	else:
		print("resource busy")

else:
	App()


